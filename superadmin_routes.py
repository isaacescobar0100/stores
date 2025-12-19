"""
Rutas del Panel SuperAdmin Profesional
"""
# Cloudflare DNS
try:
    from cloudflare_dns import crear_subdominio
except:
    crear_subdominio = None

from flask import Blueprint, render_template, request, jsonify, session, redirect, url_for, flash, g, Response
from functools import wraps
from datetime import datetime, timedelta
import csv
import io
import secrets
import string
import hashlib
from models import Tienda, Usuario, Pedido, Producto, Categoria, CategoriaMaestra, get_connection, get_placeholder

# Helper para convertir queries con ? a placeholder correcto
def q(query):
    """Convierte ? a %s para MySQL"""
    return query.replace('?', get_placeholder())
# Diccionario para traducir dias a espanol
DIAS_ESPANOL = {
    'Mon': 'Lun', 'Tue': 'Mar', 'Wed': 'Mie', 'Thu': 'Jue',
    'Fri': 'Vie', 'Sat': 'Sab', 'Sun': 'Dom'
}

def dia_espanol(fecha):
    """Convierte una fecha a nombre de dia en espanol"""
    dia_en = fecha.strftime('%a')
    return DIAS_ESPANOL.get(dia_en, dia_en)



# ============ VALIDACIÓN DE ENTRADA ============
import re

def validar_entrada(data, campos_requeridos=None, campos_opcionales=None):
    """
    Valida y sanitiza datos de entrada

    Args:
        data: dict con los datos a validar
        campos_requeridos: dict {nombre: tipo} de campos obligatorios
        campos_opcionales: dict {nombre: (tipo, default)} de campos opcionales

    Returns:
        (datos_validados, error_mensaje) - error_mensaje es None si ok
    """
    if not data or not isinstance(data, dict):
        return None, "Datos inválidos o vacíos"

    resultado = {}
    campos_requeridos = campos_requeridos or {}
    campos_opcionales = campos_opcionales or {}

    # Validar campos requeridos
    for campo, tipo in campos_requeridos.items():
        if campo not in data or data[campo] is None:
            return None, f"Campo requerido: {campo}"

        valor = data[campo]
        valor_validado = _validar_tipo(valor, tipo, campo)
        if valor_validado is None:
            return None, f"Tipo inválido para {campo}"
        resultado[campo] = valor_validado

    # Validar campos opcionales
    for campo, (tipo, default) in campos_opcionales.items():
        if campo in data and data[campo] is not None:
            valor = data[campo]
            valor_validado = _validar_tipo(valor, tipo, campo)
            if valor_validado is None:
                return None, f"Tipo inválido para {campo}"
            resultado[campo] = valor_validado
        else:
            resultado[campo] = default

    return resultado, None


def _validar_tipo(valor, tipo, campo):
    """Valida y convierte un valor al tipo esperado"""
    try:
        if tipo == 'str':
            return sanitizar_texto(str(valor).strip())
        elif tipo == 'str_slug':
            return sanitizar_slug(str(valor).strip())
        elif tipo == 'email':
            return validar_email(str(valor).strip())
        elif tipo == 'int':
            return int(valor)
        elif tipo == 'float':
            return float(valor)
        elif tipo == 'bool':
            if isinstance(valor, bool):
                return valor
            return str(valor).lower() in ('true', '1', 'yes', 'si')
        elif tipo == 'list':
            if isinstance(valor, list):
                return valor
            return None
        elif tipo == 'color':
            return validar_color(str(valor))
        else:
            return valor
    except (ValueError, TypeError):
        return None


def sanitizar_texto(texto, max_len=500):
    """Sanitiza texto eliminando caracteres peligrosos"""
    if not texto:
        return ""
    # Eliminar tags HTML
    texto = re.sub(r'<[^>]+>', '', texto)
    # Limitar longitud
    return texto[:max_len]


def sanitizar_slug(texto):
    """Sanitiza un slug/subdominio"""
    if not texto:
        return None
    # Solo letras, números y guiones
    texto = re.sub(r'[^a-zA-Z0-9-]', '', texto.lower())
    if not texto or len(texto) < 2 or len(texto) > 50:
        return None
    return texto


def validar_email(email):
    """Valida formato de email"""
    if not email:
        return None
    pattern = r'^[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}$'
    if re.match(pattern, email):
        return email.lower()
    return None


def validar_color(color):
    """Valida color hexadecimal"""
    if not color:
        return '#000000'
    if re.match(r'^#[0-9A-Fa-f]{6}$', color):
        return color
    return '#000000'



# ============ RATE LIMITING ============
from collections import defaultdict
import time

# Almacena requests por IP: {ip: [(timestamp, endpoint), ...]}
_rate_limit_store = defaultdict(list)
_RATE_LIMIT_WINDOW = 60  # segundos
_RATE_LIMIT_MAX = 60     # max requests por ventana (normal)
_RATE_LIMIT_LOGIN = 5    # max intentos login por ventana

def _limpiar_requests_viejos(ip):
    """Elimina requests fuera de la ventana de tiempo"""
    ahora = time.time()
    _rate_limit_store[ip] = [
        (ts, ep) for ts, ep in _rate_limit_store[ip]
        if ahora - ts < _RATE_LIMIT_WINDOW
    ]

def verificar_rate_limit(endpoint='general'):
    """
    Verifica si la IP ha excedido el límite de requests
    Returns: (permitido, mensaje_error)
    """
    ip = request.headers.get('CF-Connecting-IP') or request.headers.get('X-Forwarded-For', '').split(',')[0].strip() or request.remote_addr

    _limpiar_requests_viejos(ip)

    # Contar requests en la ventana actual
    requests_totales = len(_rate_limit_store[ip])
    requests_endpoint = sum(1 for ts, ep in _rate_limit_store[ip] if ep == endpoint)

    # Límites diferentes para login
    if endpoint == 'login':
        if requests_endpoint >= _RATE_LIMIT_LOGIN:
            return False, f'Demasiados intentos de login. Espera {_RATE_LIMIT_WINDOW} segundos.'
    else:
        if requests_totales >= _RATE_LIMIT_MAX:
            return False, f'Demasiadas solicitudes. Espera {_RATE_LIMIT_WINDOW} segundos.'

    # Registrar este request
    _rate_limit_store[ip].append((time.time(), endpoint))
    return True, None


def rate_limit(endpoint='general'):
    """Decorador para aplicar rate limiting"""
    def decorator(f):
        @wraps(f)
        def decorated(*args, **kwargs):
            permitido, error = verificar_rate_limit(endpoint)
            if not permitido:
                return jsonify({'error': error}), 429
            return f(*args, **kwargs)
        return decorated
    return decorator


superadmin_bp = Blueprint('superadmin', __name__, url_prefix='/superadmin')
# Desactivar cache para todas las rutas de superadmin
@superadmin_bp.after_request
def add_nocache_headers(response):
    """Agregar headers para evitar cache de Cloudflare y navegador"""
    response.headers['Cache-Control'] = 'no-store, no-cache, must-revalidate, max-age=0'
    response.headers['Pragma'] = 'no-cache'
    response.headers['Expires'] = '0'
    response.headers['CDN-Cache-Control'] = 'no-store'
    response.headers['Cloudflare-CDN-Cache-Control'] = 'no-store'
    return response




# ============ PROTECCIÓN CSRF ============

def generar_csrf_token():
    """Genera un token CSRF único para la sesión"""
    if 'csrf_token' not in session:
        session['csrf_token'] = secrets.token_hex(32)
    return session['csrf_token']


def validar_csrf_token():
    """Valida el token CSRF en requests POST/PUT/DELETE"""
    token = request.headers.get('X-CSRF-Token') or request.form.get('csrf_token') or (request.json or {}).get('csrf_token')
    if not token or token != session.get('csrf_token'):
        return False
    return True


def requiere_csrf(f):
    """Decorador: valida CSRF token en requests que modifican datos"""
    @wraps(f)
    def decorated(*args, **kwargs):
        if request.method in ['POST', 'PUT', 'DELETE']:
            if not validar_csrf_token():
                return jsonify({'error': 'Token CSRF inválido. Recarga la página.'}), 403
        return f(*args, **kwargs)
    return decorated


# Hacer el token disponible en todos los templates
@superadmin_bp.context_processor
def csrf_context():
    return {'csrf_token': generar_csrf_token}


def requiere_superadmin(f):
    """Decorador: requiere login de superadmin"""
    @wraps(f)
    def decorated(*args, **kwargs):
        if 'user_id' not in session or session.get('user_rol') != 'superadmin':
            flash('Acceso denegado', 'error')
            return redirect(url_for('superadmin.login_page'))
        return f(*args, **kwargs)
    return decorated


# ============ PÁGINAS PRINCIPALES ============

@superadmin_bp.route('/')
def login_page():
    """Página de login"""
    if 'user_id' in session and session.get('user_rol') == 'superadmin':
        return redirect(url_for('superadmin.dashboard'))
    return render_template('superadmin/login.html')


@superadmin_bp.route('/login', methods=['POST'])
@rate_limit('login')
def login():
    """Procesar login"""
    email = request.form.get('email')
    password = request.form.get('password')

    user = Usuario.validar(email, password, tienda_id=None)

    if user and user.get('rol') == 'superadmin':
        session['user_id'] = user['id']
        session['user_nombre'] = user['nombre']
        session['user_rol'] = user['rol']
        return redirect(url_for('superadmin.dashboard'))

    flash('Credenciales incorrectas o no es superadmin', 'error')
    return redirect(url_for('superadmin.login_page'))


@superadmin_bp.route('/logout')
def logout():
    """Cerrar sesión"""
    session.clear()
    return redirect(url_for('superadmin.login_page'))


@superadmin_bp.route('/dashboard')
@requiere_superadmin
def dashboard():
    """Dashboard principal"""
    db = get_connection()
    hoy = datetime.now().strftime('%Y-%m-%d')

    # Estadísticas globales
    stats = {
        'total_tiendas': db.execute('SELECT COUNT(*) FROM tiendas WHERE activo = 1').fetchone()[0],
        'total_usuarios': db.execute('SELECT COUNT(*) FROM usuarios WHERE activo = 1').fetchone()[0],
        'pedidos_hoy': db.execute(
            q('SELECT COUNT(*) FROM pedidos WHERE DATE(fecha_pedido) = ?'), (hoy,)
        ).fetchone()[0],
        'ventas_hoy': db.execute(
            q('SELECT COALESCE(SUM(total), 0) FROM pedidos WHERE DATE(fecha_pedido) = ? AND estado != "cancelado"'),
            (hoy,)
        ).fetchone()[0] or 0
    }

    # Estadísticas por tienda
    tiendas_stats = db.execute(q('''
        SELECT t.id, t.nombre, t.subdominio,
               (SELECT COUNT(*) FROM pedidos p WHERE p.tienda_id = t.id AND DATE(p.fecha_pedido) = ?) as pedidos_hoy,
               (SELECT COALESCE(SUM(p.total), 0) FROM pedidos p WHERE p.tienda_id = t.id AND DATE(p.fecha_pedido) = ? AND p.estado != 'cancelado') as ventas_hoy
        FROM tiendas t WHERE t.activo = 1
        ORDER BY ventas_hoy DESC
        LIMIT 10
    '''), (hoy, hoy)).fetchall()

    # Pedidos recientes
    pedidos_recientes = db.execute('''
        SELECT p.*, t.nombre as tienda_nombre,
               p.fecha_pedido as fecha_hora
        FROM pedidos p
        JOIN tiendas t ON p.tienda_id = t.id
        ORDER BY p.fecha_pedido DESC
        LIMIT 10
    ''').fetchall()

    # Ventas últimos 7 días
    ventas_semana = []
    max_venta = 0
    for i in range(6, -1, -1):
        fecha = (datetime.now() - timedelta(days=i)).strftime('%Y-%m-%d')
        total = db.execute(
            'SELECT COALESCE(SUM(total), 0) FROM pedidos WHERE DATE(fecha_pedido) = ? AND estado != "cancelado"',
            (fecha,)
        ).fetchone()[0] or 0
        if total > max_venta:
            max_venta = total
        ventas_semana.append({
            'fecha': fecha,
            'nombre': dia_espanol(datetime.now() - timedelta(days=i)),
            'total': total
        })

    # Calcular porcentajes
    for dia in ventas_semana:
        dia['porcentaje'] = int((dia['total'] / max_venta * 100)) if max_venta > 0 else 0

    return render_template('superadmin/dashboard.html',
                         stats=stats,
                         tiendas_stats=tiendas_stats,
                         pedidos_recientes=pedidos_recientes,
                         ventas_semana=ventas_semana)


@superadmin_bp.route('/tiendas')
@requiere_superadmin
def tiendas():
    """Gestión de tiendas"""
    db = get_connection()
    hoy = datetime.now().strftime('%Y-%m-%d')

    tiendas = db.execute(q('''
        SELECT t.*,
               (SELECT COUNT(*) FROM productos WHERE tienda_id = t.id) as total_productos,
               (SELECT COUNT(*) FROM pedidos WHERE tienda_id = t.id AND DATE(fecha_pedido) = ?) as pedidos_hoy,
               (SELECT COALESCE(SUM(total), 0) FROM pedidos WHERE tienda_id = t.id AND DATE(fecha_pedido) = ? AND estado != 'cancelado') as ventas_hoy
        FROM tiendas t
        ORDER BY t.nombre
    '''), (hoy, hoy)).fetchall()

    from flask import make_response as mr
    response = mr(render_template('superadmin/tiendas.html', tiendas=tiendas))
    response.headers['Cache-Control'] = 'no-cache, no-store, must-revalidate'
    response.headers['Pragma'] = 'no-cache'
    response.headers['Expires'] = '0'
    return response


@superadmin_bp.route('/pedidos')
@requiere_superadmin
def pedidos():
    """Ver todos los pedidos"""
    db = get_connection()

    # Filtros
    tienda_id = request.args.get('tienda')
    estado = request.args.get('estado')
    fecha = request.args.get('fecha', datetime.now().strftime('%Y-%m-%d'))

    # Query base - con JOIN a clientes para obtener nombre
    query = '''
        SELECT p.*, t.nombre as tienda_nombre,
               c.nombre as cliente_nombre, c.telefono as cliente_telefono,
               p.fecha_pedido as fecha_hora,
               (SELECT COUNT(*) FROM detalle_pedidos WHERE pedido_id = p.id) as total_items
        FROM pedidos p
        JOIN tiendas t ON p.tienda_id = t.id
        LEFT JOIN clientes c ON p.cliente_id = c.id
        WHERE 1=1
    '''
    params = []

    if tienda_id:
        query += ' AND p.tienda_id = ?'
        params.append(tienda_id)
    if estado:
        query += ' AND p.estado = ?'
        params.append(estado)
    if fecha:
        query += ' AND DATE(p.fecha_pedido) = ?'
        params.append(fecha)

    query += ' ORDER BY p.fecha_pedido DESC LIMIT 100'

    pedidos = db.execute(query, params).fetchall()
    tiendas = db.execute('SELECT id, nombre FROM tiendas WHERE activo = 1').fetchall()

    # Estadísticas
    stats = {
        'total': len(pedidos),
        'pendientes': sum(1 for p in pedidos if p['estado'] == 'pendiente'),
        'preparando': sum(1 for p in pedidos if p['estado'] == 'preparando'),
        'entregados': sum(1 for p in pedidos if p['estado'] == 'entregado'),
        'total_ventas': sum(p['total'] for p in pedidos if p['estado'] != 'cancelado')
    }

    return render_template('superadmin/pedidos.html',
                         pedidos=pedidos,
                         tiendas=tiendas,
                         stats=stats,
                         fecha_hoy=fecha)


@superadmin_bp.route('/usuarios')
@requiere_superadmin
def usuarios():
    """Gestión de usuarios"""
    db = get_connection()

    usuarios = db.execute('''
        SELECT u.*, t.nombre as tienda_nombre
        FROM usuarios u
        LEFT JOIN tiendas t ON u.tienda_id = t.id
        ORDER BY t.nombre, u.nombre
    ''').fetchall()

    tiendas = db.execute('SELECT id, nombre FROM tiendas WHERE activo = 1').fetchall()

    # Estadísticas
    stats = {
        'total': len(usuarios),
        'admins': sum(1 for u in usuarios if u['rol'] == 'admin'),
        'cocina': sum(1 for u in usuarios if u['rol'] == 'cocina'),
        'meseros': sum(1 for u in usuarios if u['rol'] == 'mesero')
    }

    return render_template('superadmin/usuarios.html',
                         usuarios=usuarios,
                         tiendas=tiendas,
                         stats=stats)


@superadmin_bp.route('/ventas')
@requiere_superadmin
def ventas():
    """Reporte de ventas"""
    db = get_connection()

    # Fechas
    fecha_hasta = request.args.get('hasta', datetime.now().strftime('%Y-%m-%d'))
    fecha_desde = request.args.get('desde', (datetime.now() - timedelta(days=30)).strftime('%Y-%m-%d'))
    tienda_id = request.args.get('tienda')

    # Query ventas por tienda
    query = '''
        SELECT t.id, t.nombre, t.subdominio,
               COUNT(p.id) as pedidos,
               COALESCE(SUM(p.total), 0) as ventas
        FROM tiendas t
        LEFT JOIN pedidos p ON t.id = p.tienda_id
            AND DATE(p.fecha_pedido) BETWEEN ? AND ?
            AND p.estado != 'cancelado'
        WHERE t.activo = 1
    '''
    params = [fecha_desde, fecha_hasta]

    if tienda_id:
        query += ' AND t.id = ?'
        params.append(tienda_id)

    query += ' GROUP BY t.id ORDER BY ventas DESC'

    ventas_por_tienda = db.execute(query, params).fetchall()

    # Calcular porcentajes
    total_ventas = sum(t['ventas'] for t in ventas_por_tienda)
    ventas_por_tienda = [dict(t) for t in ventas_por_tienda]
    for t in ventas_por_tienda:
        t['porcentaje'] = round(t['ventas'] / total_ventas * 100) if total_ventas > 0 else 0

    # Productos más vendidos
    productos_top = db.execute('''
        SELECT pr.nombre, t.nombre as tienda_nombre,
               SUM(pi.cantidad) as cantidad,
               SUM(pi.cantidad * pi.precio_unitario) as total
        FROM detalle_pedidos pi
        JOIN pedidos p ON pi.pedido_id = p.id
        JOIN productos pr ON pi.producto_id = pr.id
        JOIN tiendas t ON p.tienda_id = t.id
        WHERE DATE(p.fecha_pedido) BETWEEN ? AND ?
          AND p.estado != 'cancelado'
        GROUP BY pr.id
        ORDER BY cantidad DESC
        LIMIT 10
    ''', (fecha_desde, fecha_hasta)).fetchall()

    # Ventas diarias
    ventas_diarias = []
    max_venta = 0
    fecha_actual = datetime.strptime(fecha_desde, '%Y-%m-%d')
    fecha_fin = datetime.strptime(fecha_hasta, '%Y-%m-%d')

    while fecha_actual <= fecha_fin:
        fecha_str = fecha_actual.strftime('%Y-%m-%d')
        total = db.execute(
            'SELECT COALESCE(SUM(total), 0) FROM pedidos WHERE DATE(fecha_pedido) = ? AND estado != "cancelado"',
            (fecha_str,)
        ).fetchone()[0] or 0
        if total > max_venta:
            max_venta = total
        ventas_diarias.append({
            'fecha': fecha_str,
            'fecha_corta': fecha_actual.strftime('%d/%m'),
            'total': total
        })
        fecha_actual += timedelta(days=1)

    for dia in ventas_diarias:
        dia['porcentaje'] = int((dia['total'] / max_venta * 100)) if max_venta > 0 else 5

    # Estadísticas globales
    stats = {
        'total_ventas': total_ventas,
        'total_pedidos': sum(t['pedidos'] for t in ventas_por_tienda),
        'ticket_promedio': total_ventas / sum(t['pedidos'] for t in ventas_por_tienda) if sum(t['pedidos'] for t in ventas_por_tienda) > 0 else 0,
        'tiendas_activas': sum(1 for t in ventas_por_tienda if t['ventas'] > 0)
    }

    tiendas = db.execute('SELECT id, nombre FROM tiendas WHERE activo = 1').fetchall()

    return render_template('superadmin/ventas.html',
                         stats=stats,
                         ventas_por_tienda=ventas_por_tienda,
                         productos_top=productos_top,
                         ventas_diarias=ventas_diarias[-30:],  # Últimos 30 días max
                         tiendas=tiendas,
                         fecha_desde=fecha_desde,
                         fecha_hasta=fecha_hasta,
                         tienda_seleccionada=int(tienda_id) if tienda_id else None,
                         enumerate=enumerate)


# ============ API ENDPOINTS ============

@superadmin_bp.route('/api/tiendas/<int:id>')
@requiere_superadmin
def api_tienda_get(id):
    """Obtener datos de una tienda"""
    from decimal import Decimal
    print(f"=== API TIENDA GET id={id} ===")
    db = get_connection()
    tienda = db.execute('SELECT * FROM tiendas WHERE id = %s', (id,)).fetchone()
    print(f"Tienda encontrada: {tienda is not None}")
    if not tienda:
        return jsonify({'error': 'Tienda no encontrada'}), 404
    # Convertir Decimal a float para JSON
    result = {}
    for key, value in dict(tienda).items():
        if isinstance(value, Decimal):
            result[key] = float(value)
        else:
            result[key] = value
    # Log de campos Wompi
    print(f"Wompi activo: {result.get('wompi_activo')}")
    print(f"Wompi public: {result.get('wompi_public_key')}")
    print(f"Wompi private existe: {bool(result.get('wompi_private_key'))}")
    return jsonify(result)


@superadmin_bp.route('/api/tiendas', methods=['POST'])
@requiere_superadmin
@requiere_csrf
def api_tienda_crear():
    """Crear nueva tienda"""
    data = request.json

    # Validar entrada
    datos, error = validar_entrada(data,
        campos_requeridos={
            'nombre': 'str',
            'subdominio': 'str_slug'
        },
        campos_opcionales={
            'email': ('email', None),
            'telefono': ('str', None),
            'direccion': ('str', None),
            'horario': ('str', None),
            'slogan': ('str', None),
            'logo': ('str', None),
            'color_primario': ('color', '#ff441f'),
            'color_secundario': ('color', '#00b14f'),
            'color_terciario': ('color', '#f5f5f5'),
            'banner_url': ('str', None),
            'costo_domicilio': ('float', 0),
            'pedido_minimo': ('float', 0),
            'categoria_ids': ('list', [])
        }
    )

    if error:
        return jsonify({'error': error}), 400

    db = get_connection()

    # Verificar subdominio único
    existe = db.execute('SELECT id FROM tiendas WHERE subdominio = %s', (datos['subdominio'],)).fetchone()
    if existe:
        return jsonify({'error': 'El subdominio ya está en uso'}), 400

    try:
        cursor = db.execute(q('''
            INSERT INTO tiendas (nombre, slug, subdominio, email, telefono, direccion, horario, slogan, activo, logo, color_primario, color_secundario, color_terciario, banner_url, costo_domicilio, pedido_minimo)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, 1, ?, ?, ?, ?, ?, ?, ?)
        '''), (data['nombre'], data['subdominio'], data['subdominio'], data.get('email'), data.get('telefono'), data.get('direccion'), data.get('horario'), data.get('slogan'),
              data.get('logo'), data.get('color_primario', '#ff441f'), data.get('color_secundario', '#00b14f'), data.get('color_terciario', '#f5f5f5'), data.get('banner_url'), data.get('costo_domicilio', 0), data.get('pedido_minimo', 0)))
        db.commit()
        tienda_id = cursor.lastrowid

        # Asignar categorías si se enviaron
        categoria_ids = data.get('categoria_ids', [])
        for i, cat_id in enumerate(categoria_ids):
            CategoriaMaestra.asignar_a_tienda(tienda_id, cat_id, orden=i)

        # Crear subdominio en Cloudflare (no bloquea si falla)
        dns_resultado = None
        if crear_subdominio:
            try:
                dns_resultado = crear_subdominio(data['subdominio'])
            except Exception as dns_err:
                print(f'Error creando DNS: {dns_err}')

        # Generar contraseñas temporales
        def gen_password():
            return ''.join(secrets.choice(string.ascii_letters + string.digits) for _ in range(10))
        
        password_admin = gen_password()
        password_mesero = gen_password()
        password_cocina = gen_password()
        password_caja = gen_password()
        
        # Crear usuarios admin y cocina
        from werkzeug.security import generate_password_hash
        subdominio = data['subdominio']
        
        db.execute(q("""
            INSERT INTO usuarios (tienda_id, email, password_hash, nombre, rol)
            VALUES (?, ?, ?, ?, ?)
        """), (tienda_id, f'admin@{subdominio}.vxplay.online', generate_password_hash(password_admin), f'Admin {data["nombre"]}', 'admin'))

        db.execute(q("""
            INSERT INTO usuarios (tienda_id, email, password_hash, nombre, rol)
            VALUES (?, ?, ?, ?, ?)
        """), (tienda_id, f'mesero@{subdominio}.vxplay.online', generate_password_hash(password_mesero), f'Mesero {data["nombre"]}', 'mesero'))

        db.execute(q("""
            INSERT INTO usuarios (tienda_id, email, password_hash, nombre, rol)
            VALUES (?, ?, ?, ?, ?)
        """), (tienda_id, f'cocina@{subdominio}.vxplay.online', generate_password_hash(password_cocina), f'Cocina {data["nombre"]}', 'cocina'))

        db.execute(q("""
            INSERT INTO usuarios (tienda_id, email, password_hash, nombre, rol)
            VALUES (?, ?, ?, ?, ?)
        """), (tienda_id, f'caja@{subdominio}.vxplay.online', generate_password_hash(password_caja), f'Caja {data["nombre"]}', 'caja'))

        db.commit()

        return jsonify({
            'success': True, 
            'tienda_id': tienda_id, 
            'dns': dns_resultado,
            'credenciales': {
                'admin': {
                    'email': f'admin@{subdominio}.vxplay.online',
                    'password': password_admin
                },
                'mesero': {
                    'email': f'mesero@{subdominio}.vxplay.online',
                    'password': password_mesero
                },
                'cocina': {
                    'email': f'cocina@{subdominio}.vxplay.online',
                    'password': password_cocina
                },
                'caja': {
                    'email': f'caja@{subdominio}.vxplay.online',
                    'password': password_caja
                }
            }
        })
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@superadmin_bp.route('/api/tiendas/<int:id>', methods=['PUT'])
@requiere_superadmin
@requiere_csrf
def api_tienda_actualizar(id):
    """Actualizar tienda"""
    data = request.json

    # Log para debug
    print(f"=== ACTUALIZANDO TIENDA {id} ===")
    print(f"costo_domicilio recibido: {data.get('costo_domicilio')} (type: {type(data.get('costo_domicilio'))})")
    print(f"pedido_minimo recibido: {data.get('pedido_minimo')} (type: {type(data.get('pedido_minimo'))})")
    print(f"modo_pedido recibido: {data.get('modo_pedido')}")

    # Convertir valores a numeros correctamente
    try:
        costo_dom = float(data.get('costo_domicilio') or 0)
    except (ValueError, TypeError):
        costo_dom = 0

    try:
        pedido_min = float(data.get('pedido_minimo') or 0)
    except (ValueError, TypeError):
        pedido_min = 0

    modo_ped = data.get('modo_pedido') or 'normal'

    print(f"Valores procesados: costo_dom={costo_dom}, pedido_min={pedido_min}, modo_ped={modo_ped}")

    db = get_connection()

    try:
        db.execute(q('''
            UPDATE tiendas SET nombre = ?, subdominio = ?, email = ?, telefono = ?, direccion = ?, horario = ?, slogan = ?,
            logo = ?, color_primario = ?, color_secundario = ?, color_terciario = ?, banner_url = ?, costo_domicilio = ?, pedido_minimo = ?, modo_pedido = ?
            WHERE id = ?
        '''), (data['nombre'], data['subdominio'], data.get('email'), data.get('telefono'), data.get('direccion'), data.get('horario'), data.get('slogan'),
              data.get('logo'), data.get('color_primario'), data.get('color_secundario'), data.get('color_terciario'), data.get('banner_url'), costo_dom, pedido_min, modo_ped, id))

        # Actualizar Wompi si se envió
        wompi_activo = 1 if data.get('wompi_activo') else 0
        wompi_public = data.get('wompi_public_key', '')
        wompi_private = data.get('wompi_private_key', '')
        wompi_evento = data.get('wompi_evento_key', '')
        wompi_integrity = data.get('wompi_integrity_key', '')

        # Construir lista de campos a actualizar para Wompi
        wompi_updates = ['wompi_activo = ?', 'wompi_public_key = ?']
        wompi_values = [wompi_activo, wompi_public]

        # Solo actualizar credenciales secretas si no son asteriscos (placeholder)
        if wompi_private and wompi_private != '********':
            wompi_updates.append('wompi_private_key = ?')
            wompi_values.append(wompi_private)
        if wompi_evento and wompi_evento != '********':
            wompi_updates.append('wompi_evento_key = ?')
            wompi_values.append(wompi_evento)
        if wompi_integrity and wompi_integrity != '********':
            wompi_updates.append('wompi_integrity_key = ?')
            wompi_values.append(wompi_integrity)

        wompi_values.append(id)
        db.execute(q(f"UPDATE tiendas SET {', '.join(wompi_updates)} WHERE id = ?"), wompi_values)

        db.commit()

        # Actualizar categorías si se enviaron
        categoria_ids = data.get('categoria_ids')
        if categoria_ids is not None:
            # Obtener categorías actuales de la tienda
            cats_actuales = db.execute(q('''
                SELECT c.id, c.categoria_maestra_id
                FROM categorias c
                WHERE c.tienda_id = ?
            '''), (id,)).fetchall()

            # Crear set de IDs de categorías maestras seleccionadas
            cat_maestras_seleccionadas = set(categoria_ids)

            # Desactivar categorías no seleccionadas y activar las seleccionadas
            for cat in cats_actuales:
                cat_id = cat['id']
                cat_maestra_id = cat['categoria_maestra_id']

                if cat_maestra_id:
                    if cat_maestra_id in cat_maestras_seleccionadas:
                        # Activar si está seleccionada
                        db.execute(q('UPDATE categorias SET activo = 1 WHERE id = ?'), (cat_id,))
                    else:
                        # Desactivar si no está seleccionada
                        db.execute(q('UPDATE categorias SET activo = 0 WHERE id = ?'), (cat_id,))

            db.commit()

            # Agregar las nuevas categorías que no existen
            cats_maestras_actuales = set(c['categoria_maestra_id'] for c in cats_actuales if c['categoria_maestra_id'])
            for i, cat_id in enumerate(categoria_ids):
                if cat_id not in cats_maestras_actuales:
                    CategoriaMaestra.asignar_a_tienda(id, cat_id, orden=i)

        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@superadmin_bp.route('/api/tiendas/<int:id>/toggle', methods=['POST'])
@requiere_superadmin
@requiere_csrf
def api_tienda_toggle(id):
    """Activar/desactivar tienda"""
    db = get_connection()
    db.execute(q('UPDATE tiendas SET activo = NOT activo WHERE id = ?'), (id,))
    db.commit()
    return jsonify({'success': True})


@superadmin_bp.route('/api/tiendas/<int:id>/detalle')
@requiere_superadmin
def api_tienda_detalle(id):
    """Obtener detalle completo de tienda"""
    db = get_connection()
    hoy = datetime.now().strftime('%Y-%m-%d')

    tienda = db.execute('SELECT * FROM tiendas WHERE id = %s', (id,)).fetchone()
    if not tienda:
        return jsonify({'error': 'Tienda no encontrada'}), 404

    resultado = dict(tienda)
    resultado['total_productos'] = db.execute(q('SELECT COUNT(*) FROM productos WHERE tienda_id = ?'), (id,)).fetchone()[0]
    resultado['total_usuarios'] = db.execute(q('SELECT COUNT(*) FROM usuarios WHERE tienda_id = ?'), (id,)).fetchone()[0]
    resultado['pedidos_hoy'] = db.execute(q('SELECT COUNT(*) FROM pedidos WHERE tienda_id = ? AND DATE(fecha_pedido) = ?'), (id, hoy)).fetchone()[0]
    resultado['ventas_hoy'] = db.execute(q('SELECT COALESCE(SUM(total), 0) FROM pedidos WHERE tienda_id = ? AND DATE(fecha_pedido) = ? AND estado != "cancelado"'), (id, hoy)).fetchone()[0] or 0

    return jsonify(resultado)


@superadmin_bp.route('/api/pedidos/<int:id>')
@requiere_superadmin
def api_pedido_get(id):
    """Obtener detalle de pedido"""
    db = get_connection()

    pedido = db.execute('''
        SELECT p.*, t.nombre as tienda_nombre,
               c.nombre as cliente_nombre, c.telefono as cliente_telefono,
               p.fecha_pedido as fecha_hora
        FROM pedidos p
        JOIN tiendas t ON p.tienda_id = t.id
        LEFT JOIN clientes c ON p.cliente_id = c.id
        WHERE p.id = ?
    ''', (id,)).fetchone()

    if not pedido:
        return jsonify({'error': 'Pedido no encontrado'}), 404

    items = db.execute('''
        SELECT pi.*, pr.nombre as producto_nombre
        FROM detalle_pedidos pi
        JOIN productos pr ON pi.producto_id = pr.id
        WHERE pi.pedido_id = ?
    ''', (id,)).fetchall()

    resultado = dict(pedido)
    resultado['items'] = [dict(item) for item in items]

    return jsonify(resultado)


@superadmin_bp.route('/api/usuarios/<int:id>')
@requiere_superadmin
def api_usuario_get(id):
    """Obtener datos de usuario"""
    db = get_connection()
    usuario = db.execute(q('SELECT id, nombre, email, tienda_id, rol, activo FROM usuarios WHERE id = ?'), (id,)).fetchone()
    if not usuario:
        return jsonify({'error': 'Usuario no encontrado'}), 404
    return jsonify(dict(usuario))


@superadmin_bp.route('/api/usuarios', methods=['POST'])
@requiere_superadmin
@requiere_csrf
def api_usuario_crear():
    """Crear nuevo usuario"""
    data = request.json

    # Validar entrada
    datos, error = validar_entrada(data,
        campos_requeridos={
            'tienda_id': 'int',
            'email': 'email',
            'nombre': 'str',
            'rol': 'str',
            'password': 'str'
        }
    )
    if error:
        return jsonify({'error': error}), 400

    db = get_connection()

    # Verificar email único
    existe = db.execute(q('SELECT id FROM usuarios WHERE email = ?'), (data['email'],)).fetchone()
    if existe:
        return jsonify({'error': 'El email ya está en uso'}), 400

    try:
        from werkzeug.security import generate_password_hash
        password_hash = generate_password_hash(data['password'])

        db.execute(q('''
            INSERT INTO usuarios (nombre, email, password, tienda_id, rol, activo)
            VALUES (?, ?, ?, ?, ?, 1)
        '''), (data['nombre'], data['email'], password_hash, data['tienda_id'], data['rol']))
        db.commit()
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@superadmin_bp.route('/api/usuarios/<int:id>', methods=['PUT'])
@requiere_superadmin
@requiere_csrf
def api_usuario_actualizar(id):
    """Actualizar usuario"""
    data = request.json

    # Validar entrada
    datos, error = validar_entrada(data,
        campos_requeridos={
            'nombre': 'str',
            'rol': 'str'
        },
        campos_opcionales={
            'email': ('email', None)
        }
    )
    if error:
        return jsonify({'error': error}), 400

    db = get_connection()

    try:
        if data.get('password'):
            from werkzeug.security import generate_password_hash
            password_hash = generate_password_hash(data['password'])
            db.execute(q('''
                UPDATE usuarios SET nombre = ?, email = ?, password = ?, tienda_id = ?, rol = ?
                WHERE id = ?
            '''), (data['nombre'], data['email'], password_hash, data['tienda_id'], data['rol'], id))
        else:
            db.execute(q('''
                UPDATE usuarios SET nombre = ?, email = ?, tienda_id = ?, rol = ?
                WHERE id = ?
            '''), (data['nombre'], data['email'], data['tienda_id'], data['rol'], id))

        db.commit()
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@superadmin_bp.route('/api/usuarios/<int:id>/toggle', methods=['POST'])
@requiere_superadmin
@requiere_csrf
def api_usuario_toggle(id):
    """Activar/desactivar usuario"""
    db = get_connection()
    db.execute(q('UPDATE usuarios SET activo = NOT activo WHERE id = ?'), (id,))
    db.commit()
    return jsonify({'success': True})


@superadmin_bp.route('/api/usuarios/<int:id>/password', methods=['POST'])
@requiere_superadmin
@requiere_csrf
def api_usuario_password(id):
    """Cambiar contraseña de usuario"""
    data = request.json
    db = get_connection()

    try:
        from werkzeug.security import generate_password_hash
        password_hash = generate_password_hash(data['password'])
        db.execute(q('UPDATE usuarios SET password = ? WHERE id = ?'), (password_hash, id))
        db.commit()
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'error': str(e)}), 500


# ============ ELIMINAR TIENDA ============

@superadmin_bp.route('/api/tiendas/<int:id>', methods=['DELETE'])
@requiere_superadmin
@requiere_csrf
def api_tienda_eliminar(id):
    """Eliminar tienda y todos sus datos relacionados"""
    db = get_connection()
    cursor = db.cursor()

    try:
        # Registrar en log de auditoría
        registrar_log(db, 'eliminar_tienda', f'Eliminando tienda ID {id}')

        # Eliminar en orden para respetar foreign keys (MySQL usa %s)
        # 1. Detalle de pedidos
        db.execute('''
            DELETE FROM detalle_pedidos WHERE pedido_id IN
            (SELECT id FROM pedidos WHERE tienda_id = %s)
        ''', (id,))

        # 2. Pedidos
        db.execute('DELETE FROM pedidos WHERE tienda_id = %s', (id,))

        # 3. Clientes
        db.execute('DELETE FROM clientes WHERE tienda_id = %s', (id,))

        # 4. Productos
        db.execute('DELETE FROM productos WHERE tienda_id = %s', (id,))

        # 5. Categorías
        db.execute('DELETE FROM categorias WHERE tienda_id = %s', (id,))

        # 6. Usuarios (ANTES de tienda por foreign key)
        db.execute('DELETE FROM usuarios WHERE tienda_id = %s', (id,))

        # 7. Finalmente la tienda
        db.execute('DELETE FROM tiendas WHERE id = %s', (id,))

        db.commit()
        return jsonify({'success': True})
    except Exception as e:
        db.rollback()
        return jsonify({'error': str(e)}), 500
    finally:
        # cursor.close() no necesario
        db.close()


# ============ EXPORTAR VENTAS ============

@superadmin_bp.route('/api/ventas/exportar')
@requiere_superadmin
def api_ventas_exportar():
    """Exportar ventas a CSV/Excel"""
    db = get_connection()
    formato = request.args.get('formato', 'csv')
    fecha_desde = request.args.get('desde', (datetime.now() - timedelta(days=30)).strftime('%Y-%m-%d'))
    fecha_hasta = request.args.get('hasta', datetime.now().strftime('%Y-%m-%d'))
    tienda_id = request.args.get('tienda')

    # Query ventas
    query = '''
        SELECT p.numero_orden, t.nombre as tienda, p.total, p.estado,
               p.tipo, p.fecha_pedido, c.nombre as cliente
        FROM pedidos p
        JOIN tiendas t ON p.tienda_id = t.id
        LEFT JOIN clientes c ON p.cliente_id = c.id
        WHERE DATE(p.fecha_pedido) BETWEEN ? AND ?
    '''
    params = [fecha_desde, fecha_hasta]

    if tienda_id:
        query += ' AND p.tienda_id = ?'
        params.append(tienda_id)

    query += ' ORDER BY p.fecha_pedido DESC'
    ventas = db.execute(query, params).fetchall()

    # Generar CSV
    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow(['Orden', 'Tienda', 'Cliente', 'Total', 'Estado', 'Tipo', 'Fecha'])

    for v in ventas:
        writer.writerow([
            v['numero_orden'],
            v['tienda'],
            v['cliente'] or 'N/A',
            v['total'],
            v['estado'],
            v['tipo'],
            v['fecha_pedido']
        ])

    output.seek(0)
    filename = f'ventas_{fecha_desde}_{fecha_hasta}.csv'

    return Response(
        output.getvalue(),
        mimetype='text/csv',
        headers={'Content-Disposition': f'attachment; filename={filename}'}
    )


# ============ EXPORTAR PEDIDOS ============

@superadmin_bp.route('/api/pedidos/exportar')
@requiere_superadmin
def api_pedidos_exportar():
    """Exportar pedidos a CSV"""
    db = get_connection()
    tienda_id = request.args.get('tienda')
    estado = request.args.get('estado')
    fecha = request.args.get('fecha', datetime.now().strftime('%Y-%m-%d'))

    query = '''
        SELECT p.numero_orden, t.nombre as tienda, c.nombre as cliente,
               c.telefono, p.tipo, p.direccion_entrega, p.total, p.estado,
               p.notas, p.fecha_pedido
        FROM pedidos p
        JOIN tiendas t ON p.tienda_id = t.id
        LEFT JOIN clientes c ON p.cliente_id = c.id
        WHERE 1=1
    '''
    params = []

    if tienda_id:
        query += ' AND p.tienda_id = ?'
        params.append(tienda_id)
    if estado:
        query += ' AND p.estado = ?'
        params.append(estado)
    if fecha:
        query += ' AND DATE(p.fecha_pedido) = ?'
        params.append(fecha)

    query += ' ORDER BY p.fecha_pedido DESC'
    pedidos = db.execute(query, params).fetchall()

    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow(['Orden', 'Tienda', 'Cliente', 'Teléfono', 'Tipo', 'Dirección', 'Total', 'Estado', 'Notas', 'Fecha'])

    for p in pedidos:
        writer.writerow([
            p['numero_orden'],
            p['tienda'],
            p['cliente'] or 'N/A',
            p['telefono'] or 'N/A',
            p['tipo'],
            p['direccion_entrega'] or 'N/A',
            p['total'],
            p['estado'],
            p['notas'] or '',
            p['fecha_pedido']
        ])

    output.seek(0)
    filename = f'pedidos_{fecha}.csv'

    return Response(
        output.getvalue(),
        mimetype='text/csv',
        headers={'Content-Disposition': f'attachment; filename={filename}'}
    )


# ============ GESTIÓN DE CLIENTES ============

@superadmin_bp.route('/clientes')
@requiere_superadmin
def clientes():
    """Ver todos los clientes"""
    db = get_connection()

    clientes = db.execute('''
        SELECT c.*, t.nombre as tienda_nombre,
               (SELECT COUNT(*) FROM pedidos WHERE cliente_id = c.id) as total_pedidos,
               (SELECT COALESCE(SUM(total), 0) FROM pedidos WHERE cliente_id = c.id AND estado != 'cancelado') as total_gastado
        FROM clientes c
        JOIN tiendas t ON c.tienda_id = t.id
        ORDER BY c.fecha_registro DESC
    ''').fetchall()

    tiendas = db.execute('SELECT id, nombre FROM tiendas WHERE activo = 1').fetchall()

    stats = {
        'total': len(clientes),
        'con_pedidos': sum(1 for c in clientes if c['total_pedidos'] > 0),
        'total_gastado': sum(c['total_gastado'] for c in clientes)
    }

    return render_template('superadmin/clientes.html',
                         clientes=clientes,
                         tiendas=tiendas,
                         stats=stats)


@superadmin_bp.route('/api/clientes/<int:id>')
@requiere_superadmin
def api_cliente_get(id):
    """Obtener datos de cliente"""
    db = get_connection()
    cliente = db.execute('''
        SELECT c.*, t.nombre as tienda_nombre
        FROM clientes c
        JOIN tiendas t ON c.tienda_id = t.id
        WHERE c.id = ?
    ''', (id,)).fetchone()

    if not cliente:
        return jsonify({'error': 'Cliente no encontrado'}), 404

    # Obtener pedidos del cliente
    pedidos = db.execute('''
        SELECT id, numero_orden, total, estado, fecha_pedido
        FROM pedidos WHERE cliente_id = ?
        ORDER BY fecha_pedido DESC LIMIT 10
    ''', (id,)).fetchall()

    resultado = dict(cliente)
    resultado['pedidos'] = [dict(p) for p in pedidos]

    return jsonify(resultado)


# ============ LOGS DE AUDITORÍA ============

def registrar_log(db, accion, detalle, usuario_id=None):
    """Registrar acción en log de auditoría"""
    try:
        if usuario_id is None:
            usuario_id = session.get('user_id')
        db.execute(q('''
            INSERT INTO logs_auditoria (usuario_id, accion, detalle, fecha, ip)
            VALUES (?, ?, ?, ?, ?)
        '''), (usuario_id, accion, detalle, datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
              request.remote_addr))
    except:
        pass  # Si la tabla no existe, ignorar


@superadmin_bp.route('/logs')
@requiere_superadmin
def logs():
    """Ver logs de auditoría"""
    db = get_connection()

    # Verificar si existe la tabla
    try:
        logs = db.execute('''
            SELECT l.*, u.nombre as usuario_nombre
            FROM logs_auditoria l
            LEFT JOIN usuarios u ON l.usuario_id = u.id
            ORDER BY l.fecha DESC
            LIMIT 500
        ''').fetchall()
    except:
        logs = []

    return render_template('superadmin/logs.html', logs=logs)


# ============ PÁGINA DE EXPORTAR DATOS ============

@superadmin_bp.route('/exportar')
@requiere_superadmin
def exportar():
    """Página de exportación de datos"""
    db = get_connection()
    tiendas = db.execute('SELECT id, nombre FROM tiendas WHERE activo = 1').fetchall()
    return render_template('superadmin/exportar.html', tiendas=tiendas)


@superadmin_bp.route('/api/exportar/tienda/<int:id>')
@requiere_superadmin
def api_exportar_tienda(id):
    """Exportar todos los datos de una tienda"""
    db = get_connection()

    tienda = db.execute('SELECT * FROM tiendas WHERE id = %s', (id,)).fetchone()
    if not tienda:
        return jsonify({'error': 'Tienda no encontrada'}), 404

    output = io.StringIO()

    # Productos
    writer = csv.writer(output)
    output.write('=== PRODUCTOS ===\n')
    writer.writerow(['ID', 'Nombre', 'Precio', 'Categoría', 'Disponible'])
    productos = db.execute('''
        SELECT p.id, p.nombre, p.precio, c.nombre as categoria, p.disponible
        FROM productos p
        LEFT JOIN categorias c ON p.categoria_id = c.id
        WHERE p.tienda_id = ?
    ''', (id,)).fetchall()
    for p in productos:
        writer.writerow([p['id'], p['nombre'], p['precio'], p['categoria'] or 'Sin categoría', 'Sí' if p['disponible'] else 'No'])

    output.write('\n=== CLIENTES ===\n')
    writer.writerow(['ID', 'Nombre', 'Teléfono', 'Email', 'Dirección', 'Fecha Registro'])
    clientes = db.execute(q('SELECT * FROM clientes WHERE tienda_id = ?'), (id,)).fetchall()
    for c in clientes:
        writer.writerow([c['id'], c['nombre'], c['telefono'], c['email'] or '', c['direccion'] or '', c['fecha_registro']])

    output.write('\n=== PEDIDOS ===\n')
    writer.writerow(['Orden', 'Cliente', 'Total', 'Estado', 'Tipo', 'Fecha'])
    pedidos = db.execute('''
        SELECT p.numero_orden, c.nombre as cliente, p.total, p.estado, p.tipo, p.fecha_pedido
        FROM pedidos p
        LEFT JOIN clientes c ON p.cliente_id = c.id
        WHERE p.tienda_id = ?
        ORDER BY p.fecha_pedido DESC
    ''', (id,)).fetchall()
    for p in pedidos:
        writer.writerow([p['numero_orden'], p['cliente'] or 'N/A', p['total'], p['estado'], p['tipo'], p['fecha_pedido']])

    output.seek(0)
    filename = f'datos_tienda_{tienda["subdominio"]}_{datetime.now().strftime("%Y%m%d")}.csv'

    return Response(
        output.getvalue(),
        mimetype='text/csv',
        headers={'Content-Disposition': f'attachment; filename={filename}'}
    )


# ============ CONFIGURACIÓN DEL SISTEMA ============

@superadmin_bp.route('/configuracion')
@requiere_superadmin
def configuracion():
    """Página de configuración del sistema"""
    db = get_connection()

    # Obtener configuraciones actuales
    try:
        configs = db.execute('SELECT * FROM configuracion_sistema').fetchall()
        config_dict = {c['clave']: c['valor'] for c in configs}
    except:
        config_dict = {}

    # Estadísticas del sistema
    stats = {
        'total_tiendas': db.execute('SELECT COUNT(*) FROM tiendas').fetchone()[0],
        'total_usuarios': db.execute('SELECT COUNT(*) FROM usuarios').fetchone()[0],
        'total_pedidos': db.execute('SELECT COUNT(*) FROM pedidos').fetchone()[0],
        'total_productos': db.execute('SELECT COUNT(*) FROM productos').fetchone()[0],
        'espacio_db': 'N/A'
    }

    return render_template('superadmin/configuracion.html',
                         config=config_dict,
                         stats=stats)


@superadmin_bp.route('/api/configuracion', methods=['POST'])
@requiere_superadmin
@requiere_csrf
def api_guardar_configuracion():
    """Guardar configuración del sistema"""
    data = request.json
    db = get_connection()

    try:
        for clave, valor in data.items():
            # Verificar si existe
            existe = db.execute(q('SELECT id FROM configuracion_sistema WHERE clave = ?'), (clave,)).fetchone()
            if existe:
                db.execute(q('UPDATE configuracion_sistema SET valor = ? WHERE clave = ?'), (valor, clave))
            else:
                db.execute(q('INSERT INTO configuracion_sistema (clave, valor) VALUES (?, ?)'), (clave, valor))

        db.commit()
        registrar_log(db, 'cambiar_configuracion', f'Configuración actualizada: {list(data.keys())}')
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'error': str(e)}), 500


# ============ BÚSQUEDA GLOBAL ============

@superadmin_bp.route('/api/buscar')
@requiere_superadmin
def api_buscar():
    """Búsqueda global en tiendas, usuarios, pedidos y clientes"""
    q = request.args.get('q', '').strip()
    if len(q) < 2:
        return jsonify({'results': []})

    db = get_connection()
    results = []
    query_param = f'%{q}%'

    # Buscar tiendas
    tiendas = db.execute('''
        SELECT id, nombre, subdominio FROM tiendas
        WHERE nombre LIKE ? OR subdominio LIKE ?
        LIMIT 5
    ''', (query_param, query_param)).fetchall()

    for t in tiendas:
        results.append({
            'title': t['nombre'],
            'type': f"Tienda - {t['subdominio']}",
            'icon': 'store',
            'url': f"/superadmin/tiendas?highlight={t['id']}"
        })

    # Buscar usuarios
    usuarios = db.execute('''
        SELECT u.id, u.nombre, u.email, t.nombre as tienda
        FROM usuarios u
        LEFT JOIN tiendas t ON u.tienda_id = t.id
        WHERE u.nombre LIKE ? OR u.email LIKE ?
        LIMIT 5
    ''', (query_param, query_param)).fetchall()

    for u in usuarios:
        results.append({
            'title': u['nombre'],
            'type': f"Usuario - {u['email']}",
            'icon': 'user',
            'url': f"/superadmin/usuarios?highlight={u['id']}"
        })

    # Buscar pedidos por número de orden
    pedidos = db.execute('''
        SELECT p.id, p.numero_orden, t.nombre as tienda, c.nombre as cliente
        FROM pedidos p
        JOIN tiendas t ON p.tienda_id = t.id
        LEFT JOIN clientes c ON p.cliente_id = c.id
        WHERE p.numero_orden LIKE ?
        LIMIT 5
    ''', (query_param,)).fetchall()

    for p in pedidos:
        results.append({
            'title': f"Pedido #{p['numero_orden']}",
            'type': f"{p['tienda']} - {p['cliente'] or 'Sin cliente'}",
            'icon': 'receipt',
            'url': f"/superadmin/pedidos?orden={p['numero_orden']}"
        })

    # Buscar clientes
    clientes = db.execute('''
        SELECT c.id, c.nombre, c.telefono, t.nombre as tienda
        FROM clientes c
        JOIN tiendas t ON c.tienda_id = t.id
        WHERE c.nombre LIKE ? OR c.telefono LIKE ?
        LIMIT 5
    ''', (query_param, query_param)).fetchall()

    for c in clientes:
        results.append({
            'title': c['nombre'],
            'type': f"Cliente - {c['tienda']}",
            'icon': 'user-friends',
            'url': f"/superadmin/clientes?highlight={c['id']}"
        })

    return jsonify({'results': results[:15]})  # Máximo 15 resultados


# ============ API PAGINADA ============

@superadmin_bp.route('/api/pedidos/paginado')
@requiere_superadmin
def api_pedidos_paginado():
    """Obtener pedidos con paginación"""
    db = get_connection()
    page = int(request.args.get('page', 1))
    per_page = int(request.args.get('per_page', 20))
    tienda_id = request.args.get('tienda')
    estado = request.args.get('estado')
    fecha = request.args.get('fecha')

    offset = (page - 1) * per_page

    # Construir query
    query = '''
        SELECT p.*, t.nombre as tienda_nombre, c.nombre as cliente_nombre
        FROM pedidos p
        JOIN tiendas t ON p.tienda_id = t.id
        LEFT JOIN clientes c ON p.cliente_id = c.id
        WHERE 1=1
    '''
    count_query = 'SELECT COUNT(*) FROM pedidos p WHERE 1=1'
    params = []
    count_params = []

    if tienda_id:
        query += ' AND p.tienda_id = ?'
        count_query += ' AND p.tienda_id = ?'
        params.append(tienda_id)
        count_params.append(tienda_id)

    if estado:
        query += ' AND p.estado = ?'
        count_query += ' AND p.estado = ?'
        params.append(estado)
        count_params.append(estado)

    if fecha:
        query += ' AND DATE(p.fecha_pedido) = ?'
        count_query += ' AND DATE(p.fecha_pedido) = ?'
        params.append(fecha)
        count_params.append(fecha)

    query += ' ORDER BY p.fecha_pedido DESC LIMIT ? OFFSET ?'
    params.extend([per_page, offset])

    pedidos = db.execute(query, params).fetchall()
    total = db.execute(count_query, count_params).fetchone()[0]
    total_pages = (total + per_page - 1) // per_page

    return jsonify({
        'pedidos': [dict(p) for p in pedidos],
        'page': page,
        'per_page': per_page,
        'total': total,
        'total_pages': total_pages
    })


@superadmin_bp.route('/api/clientes/paginado')
@requiere_superadmin
def api_clientes_paginado():
    """Obtener clientes con paginación"""
    db = get_connection()
    page = int(request.args.get('page', 1))
    per_page = int(request.args.get('per_page', 20))
    tienda_id = request.args.get('tienda')
    buscar = request.args.get('buscar', '')

    offset = (page - 1) * per_page

    query = '''
        SELECT c.*, t.nombre as tienda_nombre,
               (SELECT COUNT(*) FROM pedidos WHERE cliente_id = c.id) as total_pedidos,
               (SELECT COALESCE(SUM(total), 0) FROM pedidos WHERE cliente_id = c.id AND estado != 'cancelado') as total_gastado
        FROM clientes c
        JOIN tiendas t ON c.tienda_id = t.id
        WHERE 1=1
    '''
    count_query = '''
        SELECT COUNT(*) FROM clientes c
        JOIN tiendas t ON c.tienda_id = t.id
        WHERE 1=1
    '''
    params = []
    count_params = []

    if tienda_id:
        query += ' AND c.tienda_id = ?'
        count_query += ' AND c.tienda_id = ?'
        params.append(tienda_id)
        count_params.append(tienda_id)

    if buscar:
        query += ' AND (c.nombre LIKE ? OR c.telefono LIKE ?)'
        count_query += ' AND (c.nombre LIKE ? OR c.telefono LIKE ?)'
        search_param = f'%{buscar}%'
        params.extend([search_param, search_param])
        count_params.extend([search_param, search_param])

    query += ' ORDER BY c.fecha_registro DESC LIMIT ? OFFSET ?'
    params.extend([per_page, offset])

    clientes = db.execute(q(query), params).fetchall()
    total = db.execute(q(count_query), count_params).fetchone()[0]
    total_pages = (total + per_page - 1) // per_page

    return jsonify({
        'clientes': [dict(c) for c in clientes],
        'page': page,
        'per_page': per_page,
        'total': total,
        'total_pages': total_pages
    })


# ============ GESTIÓN DE OFERTAS ============

def crear_tabla_ofertas():
    """Crear tablas de ofertas si no existen"""
    db = get_connection()
    try:
        db.execute('''
            CREATE TABLE IF NOT EXISTS ofertas (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                tienda_id INTEGER NOT NULL,
                nombre VARCHAR(100) NOT NULL,
                descripcion TEXT,
                tipo VARCHAR(20) NOT NULL DEFAULT 'descuento',
                tipo_descuento VARCHAR(20) DEFAULT 'porcentaje',
                valor_descuento DECIMAL(10,2),
                precio_combo DECIMAL(10,2),
                fecha_inicio DATE,
                fecha_fin DATE,
                activo BOOLEAN DEFAULT 1,
                fecha_creacion TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY (tienda_id) REFERENCES tiendas(id)
            )
        ''')
        db.execute('''
            CREATE TABLE IF NOT EXISTS ofertas_productos (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                oferta_id INTEGER NOT NULL,
                producto_id INTEGER NOT NULL,
                FOREIGN KEY (oferta_id) REFERENCES ofertas(id) ON DELETE CASCADE,
                FOREIGN KEY (producto_id) REFERENCES productos(id)
            )
        ''')
        db.commit()
    except:
        pass


@superadmin_bp.route('/ofertas')
@requiere_superadmin
def ofertas():
    """Gestión de ofertas y promociones"""
    db = get_connection()
    crear_tabla_ofertas()

    ofertas = db.execute('''
        SELECT o.*, t.nombre as tienda_nombre
        FROM ofertas o
        JOIN tiendas t ON o.tienda_id = t.id
        ORDER BY o.fecha_creacion DESC
    ''').fetchall()

    tiendas = db.execute('SELECT id, nombre FROM tiendas WHERE activo = 1').fetchall()

    # Estadísticas
    stats = {
        'total': len(ofertas),
        'activas': sum(1 for o in ofertas if o['activo']),
        'descuentos': sum(1 for o in ofertas if o['tipo'] == 'descuento'),
        'combos': sum(1 for o in ofertas if o['tipo'] == 'combo')
    }

    return render_template('superadmin/ofertas.html',
                         ofertas=ofertas,
                         tiendas=tiendas,
                         stats=stats)


@superadmin_bp.route('/api/ofertas/productos/<int:tienda_id>')
@requiere_superadmin
def api_ofertas_productos(tienda_id):
    """Obtener productos de una tienda para ofertas"""
    db = get_connection()
    productos = db.execute('''
        SELECT id, nombre, precio FROM productos
        WHERE tienda_id = ? AND disponible = 1
        ORDER BY nombre
    ''', (tienda_id,)).fetchall()
    return jsonify([dict(p) for p in productos])


@superadmin_bp.route('/api/ofertas/<int:id>')
@requiere_superadmin
def api_oferta_get(id):
    """Obtener datos de una oferta"""
    db = get_connection()
    oferta = db.execute(q('SELECT * FROM ofertas WHERE id = ?'), (id,)).fetchone()
    if not oferta:
        return jsonify({'error': 'Oferta no encontrada'}), 404

    resultado = dict(oferta)

    # Obtener productos asociados
    productos = db.execute('''
        SELECT producto_id FROM ofertas_productos WHERE oferta_id = ?
    ''', (id,)).fetchall()
    resultado['productos'] = [p['producto_id'] for p in productos]

    return jsonify(resultado)


@superadmin_bp.route('/api/ofertas', methods=['POST'])
@requiere_superadmin
@requiere_csrf
def api_oferta_crear():
    """Crear nueva oferta"""
    data = request.json

    # Validar entrada
    datos, error = validar_entrada(data,
        campos_requeridos={
            'tienda_id': 'int',
            'producto_id': 'int',
            'porcentaje_descuento': 'int'
        },
        campos_opcionales={
            'fecha_inicio': ('str', None),
            'fecha_fin': ('str', None)
        }
    )
    if error:
        return jsonify({'error': error}), 400

    db = get_connection()
    crear_tabla_ofertas()

    try:
        cursor = db.execute(q('''
            INSERT INTO ofertas (tienda_id, nombre, descripcion, tipo, tipo_descuento,
                                valor_descuento, precio_combo, fecha_inicio, fecha_fin, activo)
            VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, 1)
        '''), (data['tienda_id'], data['nombre'], data.get('descripcion'),
              data['tipo'], data.get('tipo_descuento', 'porcentaje'),
              data.get('valor_descuento'), data.get('precio_combo'),
              data.get('fecha_inicio'), data.get('fecha_fin')))

        oferta_id = cursor.lastrowid

        # Insertar productos asociados
        if data.get('productos'):
            for prod_id in data['productos']:
                db.execute(q('INSERT INTO ofertas_productos (oferta_id, producto_id) VALUES (?, ?)'),
                          (oferta_id, prod_id))

        db.commit()
        registrar_log(db, 'crear_oferta', f'Oferta creada: {data["nombre"]}')
        return jsonify({'success': True, 'id': oferta_id})
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@superadmin_bp.route('/api/ofertas/<int:id>', methods=['PUT'])
@requiere_superadmin
@requiere_csrf
def api_oferta_actualizar(id):
    """Actualizar oferta"""
    data = request.json
    db = get_connection()

    try:
        db.execute(q('''
            UPDATE ofertas SET tienda_id = ?, nombre = ?, descripcion = ?, tipo = ?,
                   tipo_descuento = ?, valor_descuento = ?, precio_combo = ?,
                   fecha_inicio = ?, fecha_fin = ?
            WHERE id = ?
        '''), (data['tienda_id'], data['nombre'], data.get('descripcion'),
              data['tipo'], data.get('tipo_descuento', 'porcentaje'),
              data.get('valor_descuento'), data.get('precio_combo'),
              data.get('fecha_inicio'), data.get('fecha_fin'), id))

        # Actualizar productos asociados
        db.execute(q('DELETE FROM ofertas_productos WHERE oferta_id = ?'), (id,))
        if data.get('productos'):
            for prod_id in data['productos']:
                db.execute(q('INSERT INTO ofertas_productos (oferta_id, producto_id) VALUES (?, ?)'),
                          (id, prod_id))

        db.commit()
        registrar_log(db, 'actualizar_oferta', f'Oferta actualizada: {data["nombre"]}')
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@superadmin_bp.route('/api/ofertas/<int:id>/toggle', methods=['POST'])
@requiere_superadmin
@requiere_csrf
def api_oferta_toggle(id):
    """Activar/desactivar oferta"""
    db = get_connection()
    db.execute(q('UPDATE ofertas SET activo = NOT activo WHERE id = ?'), (id,))
    db.commit()
    return jsonify({'success': True})


@superadmin_bp.route('/api/ofertas/<int:id>', methods=['DELETE'])
@requiere_superadmin
@requiere_csrf
def api_oferta_eliminar(id):
    """Eliminar oferta"""
    db = get_connection()
    try:
        db.execute(q('DELETE FROM ofertas_productos WHERE oferta_id = ?'), (id,))
        db.execute(q('DELETE FROM ofertas WHERE id = ?'), (id,))
        db.commit()
        registrar_log(db, 'eliminar_oferta', f'Oferta eliminada ID: {id}')
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'error': str(e)}), 500


# ============ GESTIÓN DE CATEGORÍAS MAESTRAS ============

def _convertir_iconos_absolutos(categorias):
    """Convierte URLs de iconos relativos a absolutos"""
    host = request.host
    for cat in categorias:
        if cat.get('icono_url') and cat['icono_url'].startswith('/static/'):
            if 'vxplay.online' in host and not host.startswith('vxplay.online'):
                cat['icono_url'] = f"https://{host}{cat['icono_url']}"
            else:
                cat['icono_url'] = f"https://guacherna.vxplay.online{cat['icono_url']}"
    return categorias


@superadmin_bp.route('/categorias')
@requiere_superadmin
def categorias_maestras():
    """Gestión de categorías maestras"""
    categorias = CategoriaMaestra.obtener_todas(solo_activas=False)
    tiendas = Tienda.obtener_todas()

    stats = {
        'total': len(categorias),
        'activas': sum(1 for c in categorias if c['activo'])
    }

    return render_template('superadmin/categorias.html',
                         categorias=categorias,
                         tiendas=tiendas,
                         stats=stats)


@superadmin_bp.route('/api/categorias-maestras')
@requiere_superadmin
def api_categorias_maestras():
    """Obtener todas las categorías maestras activas"""
    categorias = CategoriaMaestra.obtener_todas(solo_activas=True)
    categorias = _convertir_iconos_absolutos(categorias)
    return jsonify(categorias)


@superadmin_bp.route('/api/categorias-maestras/<int:id>')
@requiere_superadmin
def api_categoria_maestra_get(id):
    """Obtener una categoría maestra"""
    cat = CategoriaMaestra.obtener_por_id(id)
    if not cat:
        return jsonify({'error': 'Categoría no encontrada'}), 404
    return jsonify(cat)


@superadmin_bp.route('/api/categorias-maestras', methods=['POST'])
@requiere_superadmin
@requiere_csrf
def api_categoria_maestra_crear():
    """Crear nueva categoría maestra"""
    import base64
    import uuid
    import os

    data = request.json
    icono_url = data.get('icono_url', '')

    # Si viene icono en base64, guardarlo como archivo
    if 'icono_base64' in data and data['icono_base64']:
        try:
            base64_data = data['icono_base64']
            # Extraer tipo y datos
            if ',' in base64_data:
                header, base64_content = base64_data.split(',', 1)
                # Determinar extension
                if 'png' in header:
                    ext = 'png'
                elif 'jpeg' in header or 'jpg' in header:
                    ext = 'jpg'
                elif 'gif' in header:
                    ext = 'gif'
                elif 'webp' in header:
                    ext = 'webp'
                else:
                    ext = 'png'
            else:
                base64_content = base64_data
                ext = 'png'

            # Decodificar y guardar
            img_bytes = base64.b64decode(base64_content)
            filename = f"cat_{uuid.uuid4().hex[:8]}.{ext}"
            upload_dir = os.path.join(os.path.dirname(__file__), 'static', 'uploads', 'categorias')
            os.makedirs(upload_dir, exist_ok=True)
            filepath = os.path.join(upload_dir, filename)

            with open(filepath, 'wb') as f:
                f.write(img_bytes)

            icono_url = f'/static/uploads/categorias/{filename}'
        except Exception as e:
            print(f"Error guardando icono: {e}")
            return jsonify({'error': 'Error al guardar el icono'}), 500

    cat_id = CategoriaMaestra.crear(data['nombre'], icono_url, data.get('orden', 0))
    if cat_id:
        return jsonify({'success': True, 'id': cat_id})
    return jsonify({'error': 'Ya existe una categoría con ese nombre'}), 400




@superadmin_bp.route('/api/upload-categoria-icon', methods=['POST'])
@requiere_superadmin
def api_upload_categoria_icon():
    """Subir icono de categoria"""
    if 'file' not in request.files:
        return jsonify({'error': 'No se envio archivo'}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'Archivo vacio'}), 400

    if file:
        import uuid
        import os
        ext = file.filename.rsplit('.', 1)[-1].lower() if '.' in file.filename else 'png'
        if ext not in ['png', 'jpg', 'jpeg', 'gif', 'webp']:
            return jsonify({'error': 'Formato no permitido'}), 400

        filename = f"cat_{uuid.uuid4().hex[:8]}.{ext}"
        upload_path = os.path.join(os.path.dirname(__file__), 'static', 'uploads', filename)
        os.makedirs(os.path.dirname(upload_path), exist_ok=True)
        file.save(upload_path)

        return jsonify({'success': True, 'url': f'/static/uploads/{filename}'})

    return jsonify({'error': 'Error al subir archivo'}), 500

@superadmin_bp.route('/api/categorias-maestras/<int:id>', methods=['PUT'])
@requiere_superadmin
@requiere_csrf
def api_categoria_maestra_actualizar(id):
    """Actualizar categoría maestra"""
    import base64
    import uuid
    import os

    data = request.json

    # Si viene icono en base64, guardarlo como archivo
    if 'icono_base64' in data and data['icono_base64']:
        try:
            base64_data = data['icono_base64']
            # Extraer tipo y datos
            if ',' in base64_data:
                header, base64_content = base64_data.split(',', 1)
                # Determinar extension
                if 'png' in header:
                    ext = 'png'
                elif 'jpeg' in header or 'jpg' in header:
                    ext = 'jpg'
                elif 'gif' in header:
                    ext = 'gif'
                elif 'webp' in header:
                    ext = 'webp'
                else:
                    ext = 'png'
            else:
                base64_content = base64_data
                ext = 'png'

            # Decodificar y guardar
            img_bytes = base64.b64decode(base64_content)
            filename = f"cat_{uuid.uuid4().hex[:8]}.{ext}"
            upload_dir = os.path.join(os.path.dirname(__file__), 'static', 'uploads', 'categorias')
            os.makedirs(upload_dir, exist_ok=True)
            filepath = os.path.join(upload_dir, filename)

            with open(filepath, 'wb') as f:
                f.write(img_bytes)

            # Usar la nueva URL
            data['icono_url'] = f'/static/uploads/categorias/{filename}'
            del data['icono_base64']
        except Exception as e:
            print(f"Error guardando icono: {e}")
            return jsonify({'error': 'Error al guardar el icono'}), 500

    CategoriaMaestra.actualizar(id, **data)

    # Sincronizar icono con categorias de tiendas que usan esta maestra
    if 'icono_url' in data:
        db = get_connection()
        db.execute(
            "UPDATE categorias SET icono = %s WHERE categoria_maestra_id = %s",
            (data['icono_url'], id)
        )
        db.commit()

    return jsonify({'success': True})


@superadmin_bp.route('/api/categorias-maestras/<int:id>/toggle', methods=['POST'])
@requiere_superadmin
@requiere_csrf
def api_categoria_maestra_toggle(id):
    """Activar/desactivar categoría maestra"""
    db = get_connection()
    db.execute(q('UPDATE categorias_maestras SET activo = NOT activo WHERE id = ?'), (id,))
    db.commit()
    return jsonify({'success': True})


@superadmin_bp.route('/api/categorias-maestras/<int:id>', methods=['DELETE'])
@requiere_superadmin
@requiere_csrf
def api_categoria_maestra_eliminar(id):
    """Eliminar categoría maestra (soft delete)"""
    CategoriaMaestra.eliminar(id)
    return jsonify({'success': True})


# ============ ASIGNAR CATEGORÍAS A TIENDAS ============

@superadmin_bp.route('/tiendas/<int:tienda_id>/categorias')
@requiere_superadmin
def tienda_categorias(tienda_id):
    """Ver y asignar categorías a una tienda"""
    tienda = Tienda.obtener_por_id(tienda_id)
    if not tienda:
        flash('Tienda no encontrada', 'error')
        return redirect(url_for('superadmin.tiendas'))

    categorias = CategoriaMaestra.obtener_para_tienda(tienda_id)
    categorias = _convertir_iconos_absolutos(categorias)

    return render_template('superadmin/tienda_categorias.html',
                         tienda=tienda,
                         categorias=categorias)


@superadmin_bp.route('/api/tiendas/<int:tienda_id>/categorias')
@requiere_superadmin
def api_tienda_categorias(tienda_id):
    """Obtener categorías de una tienda con estado de asignación"""
    categorias = CategoriaMaestra.obtener_para_tienda(tienda_id)
    categorias = _convertir_iconos_absolutos(categorias)
    return jsonify(categorias)


@superadmin_bp.route('/api/tiendas/<int:tienda_id>/categorias/asignar', methods=['POST'])
@requiere_superadmin
@requiere_csrf
def api_asignar_categoria(tienda_id):
    """Asignar categoría maestra a tienda"""
    data = request.json
    categoria_maestra_id = data.get('categoria_maestra_id')

    result = CategoriaMaestra.asignar_a_tienda(tienda_id, categoria_maestra_id)
    if result:
        return jsonify({'success': True, 'categoria_id': result})
    return jsonify({'error': 'No se pudo asignar la categoría'}), 400


@superadmin_bp.route('/api/tiendas/<int:tienda_id>/categorias/quitar', methods=['POST'])
@requiere_superadmin
@requiere_csrf
def api_quitar_categoria(tienda_id):
    """Quitar categoría de tienda"""
    data = request.json
    categoria_id = data.get('categoria_id')

    db = get_connection()
    p = get_placeholder()
    db.execute(f'DELETE FROM categorias WHERE id = {p} AND tienda_id = {p}', (categoria_id, tienda_id))
    db.commit()
    db.close()
    return jsonify({'success': True})


@superadmin_bp.route('/api/tiendas/<int:tienda_id>/categorias/asignar-multiples', methods=['POST'])
@requiere_superadmin
@requiere_csrf
def api_asignar_multiples_categorias(tienda_id):
    """Asignar múltiples categorías maestras a tienda"""
    data = request.json
    categoria_ids = data.get('categoria_ids', [])

    asignadas = 0
    for i, cat_id in enumerate(categoria_ids):
        result = CategoriaMaestra.asignar_a_tienda(tienda_id, cat_id, orden=i)
        if result:
            asignadas += 1

    return jsonify({'success': True, 'asignadas': asignadas})


@superadmin_bp.route('/tiendas/<int:tienda_id>/usuarios')
@requiere_superadmin
def tienda_usuarios(tienda_id):
    """Ver usuarios de una tienda específica"""
    db = get_connection()
    
    tienda = db.execute(q('SELECT * FROM tiendas WHERE id = ?'), (tienda_id,)).fetchone()
    if not tienda:
        return render_template('superadmin/error.html', mensaje='Tienda no encontrada'), 404

    usuarios = db.execute(q('''
        SELECT * FROM usuarios
        WHERE tienda_id = ?
        ORDER BY nombre
    '''), (tienda_id,)).fetchall()
    
    return render_template('superadmin/tienda_usuarios.html',
                         tienda=tienda,
                         usuarios=usuarios)



@superadmin_bp.route('/descargar-plantilla-word')
@requiere_superadmin
def descargar_plantilla_word():
    """Generar y descargar plantilla Word para importar productos"""
    from docx import Document
    from docx.shared import Cm
    from io import BytesIO

    doc = Document()

    # Titulo
    doc.add_heading('Plantilla de Importacion de Productos', 0)

    # Instrucciones
    doc.add_paragraph('Instrucciones:', style='Heading 2')
    instrucciones = doc.add_paragraph()
    instrucciones.add_run('1. Llena cada fila con los datos del producto\n')
    instrucciones.add_run('2. En la columna "Imagen", pega la foto directamente (clic derecho > Pegar)\n')
    instrucciones.add_run('3. La categoria debe coincidir exactamente con las categorias de tu tienda\n')
    instrucciones.add_run('4. Subcategoria y Descripcion son opcionales\n')
    instrucciones.add_run('5. Guarda el documento y subelo en el sistema\n')

    doc.add_paragraph()

    # Crear tabla
    table = doc.add_table(rows=4, cols=6)
    table.style = 'Table Grid'

    # Encabezados
    headers = ['Nombre', 'Precio', 'Categoria', 'Subcategoria', 'Descripcion', 'Imagen']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    # Filas de ejemplo
    ejemplos = [
        ['Hamburguesa Clasica', '15000', 'Hamburguesas', '', 'Deliciosa hamburguesa', '[Pegar imagen]'],
        ['Pizza Pepperoni', '25000', 'Pizzas', 'Grandes', '', '[Pegar imagen]'],
        ['Coca Cola 400ml', '5000', 'Bebidas', 'Gaseosas', '', '[Pegar imagen]'],
    ]

    for row_idx, ejemplo in enumerate(ejemplos):
        for col_idx, valor in enumerate(ejemplo):
            table.rows[row_idx + 1].cells[col_idx].text = valor

    # Guardar en memoria
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return Response(
        buffer.getvalue(),
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        headers={'Content-Disposition': 'attachment; filename=plantilla_productos.docx'}
    )



@superadmin_bp.route('/api/tienda/<int:tienda_id>/categorias')
@requiere_superadmin
def api_obtener_categorias_tienda(tienda_id):
    """Obtener categorias de una tienda desde tabla categorias"""
    db = get_connection()
    categorias = db.execute(
        "SELECT id, nombre, padre_id FROM categorias WHERE tienda_id = %s ORDER BY nombre",
        (tienda_id,)
    ).fetchall()
    return jsonify({'categorias': [dict(c) for c in categorias]})

@superadmin_bp.route('/importar-productos')
@requiere_superadmin
def importar_productos():
    """Pagina para importar productos desde Excel/CSV"""
    db = get_connection()
    tiendas = db.execute("SELECT id, nombre FROM tiendas WHERE activo = 1 ORDER BY nombre").fetchall()
    return render_template('superadmin/importar_productos.html', tiendas=tiendas)


@superadmin_bp.route('/galeria-productos')
@requiere_superadmin
def galeria_productos():
    """Pagina para gestionar imagenes de productos"""
    db = get_connection()
    tiendas = db.execute("SELECT id, nombre FROM tiendas WHERE activo = 1 ORDER BY nombre").fetchall()
    tienda_seleccionada = request.args.get('tienda', type=int)
    return render_template('superadmin/galeria_productos.html',
                         tiendas=tiendas,
                         tienda_seleccionada=tienda_seleccionada)


@superadmin_bp.route('/api/galeria/productos')
@requiere_superadmin
def api_galeria_productos():
    """Obtener productos para la galeria con filtros"""
    tienda_id = request.args.get('tienda_id', type=int)
    categoria_id = request.args.get('categoria_id', type=int)
    filtro_imagen = request.args.get('filtro_imagen', '')

    if not tienda_id:
        return jsonify({'error': 'tienda_id requerido'}), 400

    db = get_connection()

    query = """
        SELECT p.id, p.nombre, p.precio, p.imagen, p.disponible,
               c.nombre as categoria_nombre
        FROM productos p
        LEFT JOIN categorias c ON p.categoria_id = c.id
        WHERE p.tienda_id = %s
    """
    params = [tienda_id]

    if categoria_id:
        query += " AND p.categoria_id = %s"
        params.append(categoria_id)

    if filtro_imagen == 'sin_imagen':
        query += " AND (p.imagen IS NULL OR p.imagen = '' OR p.imagen LIKE '%%default%%')"
    elif filtro_imagen == 'con_imagen':
        query += " AND p.imagen IS NOT NULL AND p.imagen != '' AND p.imagen NOT LIKE '%%default%%'"

    query += " ORDER BY p.nombre"

    productos = db.execute(query, params).fetchall()

    # Calcular estadisticas
    total = len(productos)
    con_imagen = sum(1 for p in productos if p['imagen'] and 'default' not in (p['imagen'] or ''))
    sin_imagen = total - con_imagen

    return jsonify({
        'productos': [dict(p) for p in productos],
        'total': total,
        'con_imagen': con_imagen,
        'sin_imagen': sin_imagen
    })


@superadmin_bp.route('/api/galeria/producto/imagen', methods=['POST'])
@requiere_superadmin
def api_galeria_actualizar_imagen():
    """Actualizar imagen de un producto"""
    import os
    import uuid
    import base64

    try:
        data = request.get_json()
        producto_id = data.get('producto_id')
        imagen_base64 = data.get('imagen_base64')

        if not producto_id or not imagen_base64:
            return jsonify({'error': 'Datos incompletos'}), 400

        db = get_connection()

        # Verificar que el producto existe
        producto = db.execute("SELECT id, tienda_id, imagen FROM productos WHERE id = %s", (producto_id,)).fetchone()
        if not producto:
            return jsonify({'error': 'Producto no encontrado'}), 404

        # Directorio para imagenes
        upload_dir = os.path.join(os.path.dirname(__file__), 'static', 'uploads', 'productos')
        os.makedirs(upload_dir, exist_ok=True)

        # Procesar imagen base64
        if ',' in imagen_base64:
            header, img_data = imagen_base64.split(',', 1)
            ext = 'png'
            if 'jpeg' in header or 'jpg' in header:
                ext = 'jpg'
            elif 'gif' in header:
                ext = 'gif'
            elif 'webp' in header:
                ext = 'webp'
        else:
            img_data = imagen_base64
            ext = 'png'

        # Guardar imagen
        filename = f"prod_{producto['tienda_id']}_{uuid.uuid4().hex[:8]}.{ext}"
        filepath = os.path.join(upload_dir, filename)
        with open(filepath, 'wb') as f:
            f.write(base64.b64decode(img_data))

        imagen_url = f'/static/uploads/productos/{filename}'

        # Eliminar imagen anterior si existe
        if producto['imagen'] and '/uploads/' in (producto['imagen'] or ''):
            old_path = os.path.join(os.path.dirname(__file__), producto['imagen'].lstrip('/'))
            if os.path.exists(old_path):
                try:
                    os.remove(old_path)
                except:
                    pass

        # Actualizar en BD
        db.execute("UPDATE productos SET imagen = %s WHERE id = %s", (imagen_url, producto_id))
        db.commit()

        return jsonify({'success': True, 'imagen_url': imagen_url})

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500



@superadmin_bp.route('/api/importar-productos', methods=['POST'])
@requiere_superadmin
def api_importar_productos():
    """Procesar importacion masiva de productos con imagenes"""
    import os
    import uuid
    import base64

    try:
        data = request.get_json()
        tienda_id = data.get('tienda_id')
        productos = data.get('productos', [])

        if not tienda_id or not productos:
            return jsonify({'error': 'Datos incompletos'}), 400

        db = get_connection()

        # Cargar todas las categorias de la tienda
        categorias_tienda = db.execute(
            "SELECT id, nombre, padre_id FROM categorias WHERE tienda_id = %s",
            (tienda_id,)
        ).fetchall()

        # Cache de categorias (insensible a mayusculas)
        categorias_cache = {}
        for cat in categorias_tienda:
            key = cat['nombre'].lower().strip()
            if cat['padre_id'] is None:
                categorias_cache[key] = {'id': cat['id'], 'nombre': cat['nombre']}
            else:
                padre = next((c for c in categorias_tienda if c['id'] == cat['padre_id']), None)
                if padre:
                    subkey = f"{padre['nombre'].lower().strip()}_{key}"
                    categorias_cache[subkey] = {'id': cat['id'], 'nombre': cat['nombre'], 'padre_id': cat['padre_id']}

        # Directorio para imagenes
        upload_dir = os.path.join(os.path.dirname(__file__), 'static', 'uploads', 'productos')
        os.makedirs(upload_dir, exist_ok=True)

        importados = 0
        errores = 0
        errores_detalle = []

        for idx, prod in enumerate(productos):
            try:
                nombre = (prod.get('nombre') or prod.get('Nombre') or '').strip()
                precio = prod.get('precio') or prod.get('Precio', 0)
                categoria_nombre = (prod.get('categoria') or prod.get('Categoria') or '').strip()
                subcategoria_nombre = (prod.get('subcategoria') or prod.get('Subcategoria') or '').strip()
                descripcion = (prod.get('descripcion') or prod.get('Descripcion') or '').strip()
                imagen_base64 = prod.get('imagen_base64')
                imagen_nombre = prod.get('imagen_nombre', '')

                if not nombre:
                    errores += 1
                    errores_detalle.append(f"Fila {idx+1}: Sin nombre")
                    continue

                # Convertir precio
                try:
                    precio = float(str(precio).replace(',', '').replace('$', '').strip())
                except:
                    precio = 0

                # Buscar categoria
                cat_key = categoria_nombre.lower().strip()
                categoria_id = None

                if cat_key in categorias_cache:
                    padre_cat = categorias_cache[cat_key]
                    if subcategoria_nombre:
                        subcat_key = f"{cat_key}_{subcategoria_nombre.lower().strip()}"
                        if subcat_key in categorias_cache:
                            categoria_id = categorias_cache[subcat_key]['id']
                        else:
                            categoria_id = padre_cat['id']
                    else:
                        categoria_id = padre_cat['id']
                else:
                    errores += 1
                    errores_detalle.append(f"Fila {idx+1}: Categoria '{categoria_nombre}' no existe")
                    continue

                # Procesar imagen
                imagen_url = '/static/img/producto_default.png'
                if imagen_base64:
                    try:
                        # Extraer extension y datos
                        if ',' in imagen_base64:
                            header, img_data = imagen_base64.split(',', 1)
                            ext = 'png'
                            if 'jpeg' in header or 'jpg' in header:
                                ext = 'jpg'
                            elif 'gif' in header:
                                ext = 'gif'
                            elif 'webp' in header:
                                ext = 'webp'
                        else:
                            img_data = imagen_base64
                            ext = imagen_nombre.split('.')[-1] if '.' in imagen_nombre else 'png'

                        # Guardar imagen
                        filename = f"prod_{tienda_id}_{uuid.uuid4().hex[:8]}.{ext}"
                        filepath = os.path.join(upload_dir, filename)
                        with open(filepath, 'wb') as f:
                            f.write(base64.b64decode(img_data))
                        imagen_url = f'/static/uploads/productos/{filename}'
                    except Exception as img_err:
                        errores_detalle.append(f"Fila {idx+1}: Error al guardar imagen - {str(img_err)}")

                # Insertar producto
                db.execute(
                    """INSERT INTO productos (tienda_id, categoria_id, nombre, descripcion, precio, disponible, imagen)
                       VALUES (%s, %s, %s, %s, %s, 1, %s)""",
                    (tienda_id, categoria_id, nombre, descripcion, precio, imagen_url)
                )
                importados += 1

            except Exception as e:
                errores += 1
                errores_detalle.append(f"Fila {idx+1}: Error - {str(e)}")

        db.commit()

        return jsonify({
            'success': True,
            'importados': importados,
            'errores': errores,
            'errores_detalle': errores_detalle[:10]
        })

    except Exception as e:
        import traceback
        traceback.print_exc()
        return jsonify({'error': str(e)}), 500



# ============ REPORTES AVANZADOS ============

@superadmin_bp.route('/reportes/')
@requiere_superadmin
def reportes_index():
    """Vista principal de reportes avanzados"""
    db = get_connection()

    fecha_hasta = request.args.get('hasta', datetime.now().strftime('%Y-%m-%d'))
    fecha_desde = request.args.get('desde', (datetime.now() - timedelta(days=30)).strftime('%Y-%m-%d'))
    tienda_id = request.args.get('tienda')

    tiendas = db.execute('SELECT id, nombre FROM tiendas WHERE activo = 1 ORDER BY nombre').fetchall()
    db.close()

    return render_template('superadmin/reportes_avanzados.html',
                         tiendas=tiendas,
                         fecha_desde=fecha_desde,
                         fecha_hasta=fecha_hasta,
                         tienda_seleccionada=int(tienda_id) if tienda_id else None)


@superadmin_bp.route('/reportes/exportar/excel')
@requiere_superadmin
def reportes_exportar_excel():
    """Exportar reporte completo a Excel con multiples hojas"""
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill
    from openpyxl.utils import get_column_letter

    db = get_connection()
    fecha_desde = request.args.get('desde', (datetime.now() - timedelta(days=30)).strftime('%Y-%m-%d'))
    fecha_hasta = request.args.get('hasta', datetime.now().strftime('%Y-%m-%d'))
    tienda_id = request.args.get('tienda')

    wb = Workbook()
    header_font = Font(bold=True, color='FFFFFF')
    header_fill = PatternFill(start_color='FF6B35', end_color='FF6B35', fill_type='solid')

    # HOJA 1: RESUMEN
    ws1 = wb.active
    ws1.title = "Resumen"

    query = """SELECT COUNT(p.id) as pedidos, COALESCE(SUM(p.total), 0) as ventas,
               COUNT(DISTINCT p.cliente_id) as clientes
               FROM pedidos p WHERE DATE(p.fecha_pedido) BETWEEN %s AND %s AND p.estado != 'cancelado'"""
    params = [fecha_desde, fecha_hasta]
    if tienda_id:
        query += ' AND p.tienda_id = %s'
        params.append(tienda_id)

    stats = db.execute(query, params).fetchone()

    ws1['A1'] = 'REPORTE DE VENTAS'
    ws1['A1'].font = Font(bold=True, size=16)
    ws1['A2'] = f'Periodo: {fecha_desde} a {fecha_hasta}'
    ws1['A4'] = 'Metrica'; ws1['B4'] = 'Valor'
    ws1['A4'].font = header_font; ws1['A4'].fill = header_fill
    ws1['B4'].font = header_font; ws1['B4'].fill = header_fill
    ws1['A5'] = 'Total Ventas'; ws1['B5'] = float(stats['ventas'])
    ws1['A6'] = 'Total Pedidos'; ws1['B6'] = stats['pedidos']
    ws1['A7'] = 'Clientes Unicos'; ws1['B7'] = stats['clientes']

    # HOJA 2: VENTAS POR TIENDA
    ws2 = wb.create_sheet("Ventas por Tienda")
    query_t = """SELECT t.nombre, COUNT(p.id) as pedidos, COALESCE(SUM(p.total), 0) as ventas
                 FROM tiendas t LEFT JOIN pedidos p ON t.id = p.tienda_id
                 AND DATE(p.fecha_pedido) BETWEEN %s AND %s AND p.estado != 'cancelado'
                 WHERE t.activo = 1"""
    params_t = [fecha_desde, fecha_hasta]
    if tienda_id:
        query_t += ' AND t.id = %s'
        params_t.append(tienda_id)
    query_t += ' GROUP BY t.id ORDER BY ventas DESC'

    tiendas_data = db.execute(query_t, params_t).fetchall()

    for col, h in enumerate(['Tienda', 'Pedidos', 'Ventas'], 1):
        c = ws2.cell(row=1, column=col, value=h)
        c.font = header_font; c.fill = header_fill

    for row, t in enumerate(tiendas_data, 2):
        ws2.cell(row=row, column=1, value=t['nombre'])
        ws2.cell(row=row, column=2, value=t['pedidos'])
        ws2.cell(row=row, column=3, value=float(t['ventas']))

    # HOJA 3: VENTAS POR HORA
    ws3 = wb.create_sheet("Ventas por Hora")
    query_h = """SELECT HOUR(fecha_pedido) as hora, COUNT(*) as pedidos, SUM(total) as ventas
                 FROM pedidos WHERE DATE(fecha_pedido) BETWEEN %s AND %s AND estado != 'cancelado'"""
    params_h = [fecha_desde, fecha_hasta]
    if tienda_id:
        query_h += ' AND tienda_id = %s'
        params_h.append(tienda_id)
    query_h += ' GROUP BY HOUR(fecha_pedido) ORDER BY hora'

    horas_data = db.execute(query_h, params_h).fetchall()

    for col, h in enumerate(['Hora', 'Pedidos', 'Ventas'], 1):
        c = ws3.cell(row=1, column=col, value=h)
        c.font = header_font; c.fill = header_fill

    for row, h in enumerate(horas_data, 2):
        ws3.cell(row=row, column=1, value=f"{h['hora']:02d}:00")
        ws3.cell(row=row, column=2, value=h['pedidos'])
        ws3.cell(row=row, column=3, value=float(h['ventas']))

    # HOJA 4: PRODUCTOS TOP
    ws4 = wb.create_sheet("Top Productos")
    query_p = """SELECT pr.nombre, t.nombre as tienda, SUM(dp.cantidad) as cantidad,
                 SUM(dp.cantidad * dp.precio_unitario) as total
                 FROM detalle_pedidos dp JOIN pedidos p ON dp.pedido_id = p.id
                 JOIN productos pr ON dp.producto_id = pr.id JOIN tiendas t ON p.tienda_id = t.id
                 WHERE DATE(p.fecha_pedido) BETWEEN %s AND %s AND p.estado != 'cancelado'"""
    params_p = [fecha_desde, fecha_hasta]
    if tienda_id:
        query_p += ' AND p.tienda_id = %s'
        params_p.append(tienda_id)
    query_p += ' GROUP BY pr.id ORDER BY cantidad DESC LIMIT 50'

    productos = db.execute(query_p, params_p).fetchall()

    for col, h in enumerate(['Producto', 'Tienda', 'Cantidad', 'Total'], 1):
        c = ws4.cell(row=1, column=col, value=h)
        c.font = header_font; c.fill = header_fill

    for row, p in enumerate(productos, 2):
        ws4.cell(row=row, column=1, value=p['nombre'])
        ws4.cell(row=row, column=2, value=p['tienda'])
        ws4.cell(row=row, column=3, value=int(p['cantidad']))
        ws4.cell(row=row, column=4, value=float(p['total']))

    db.close()

    output = io.BytesIO()
    wb.save(output)
    output.seek(0)

    filename = f'reporte_ventas_{fecha_desde}_{fecha_hasta}.xlsx'
    return Response(output.getvalue(),
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        headers={'Content-Disposition': f'attachment; filename={filename}'})


@superadmin_bp.route('/reportes/api/ventas-por-hora')
@requiere_superadmin
def api_ventas_por_hora():
    """Ventas agrupadas por hora del dia"""
    db = get_connection()
    fecha_desde = request.args.get('desde', (datetime.now() - timedelta(days=30)).strftime('%Y-%m-%d'))
    fecha_hasta = request.args.get('hasta', datetime.now().strftime('%Y-%m-%d'))
    tienda_id = request.args.get('tienda')

    query = """SELECT HOUR(fecha_pedido) as hora, COUNT(*) as pedidos, SUM(total) as ventas
               FROM pedidos WHERE DATE(fecha_pedido) BETWEEN %s AND %s AND estado != 'cancelado'"""
    params = [fecha_desde, fecha_hasta]
    if tienda_id:
        query += ' AND tienda_id = %s'
        params.append(tienda_id)
    query += ' GROUP BY HOUR(fecha_pedido) ORDER BY hora'

    resultados = db.execute(query, params).fetchall()
    db.close()

    horas_dict = {r['hora']: r for r in resultados}
    horas = []
    for h in range(24):
        if h in horas_dict:
            horas.append({'hora': h, 'label': f'{h:02d}:00', 'pedidos': horas_dict[h]['pedidos'], 'ventas': float(horas_dict[h]['ventas'])})
        else:
            horas.append({'hora': h, 'label': f'{h:02d}:00', 'pedidos': 0, 'ventas': 0})

    return jsonify(horas)


@superadmin_bp.route('/reportes/api/comparativo-mensual')
@requiere_superadmin
def api_comparativo_mensual():
    """Comparar mes actual vs mes anterior"""
    db = get_connection()
    tienda_id = request.args.get('tienda')

    hoy = datetime.now()
    inicio_mes_actual = hoy.replace(day=1)
    fin_mes_anterior = inicio_mes_actual - timedelta(days=1)
    inicio_mes_anterior = fin_mes_anterior.replace(day=1)

    def get_stats(inicio, fin):
        query = """SELECT COUNT(*) as pedidos, COALESCE(SUM(total), 0) as ventas, COUNT(DISTINCT cliente_id) as clientes
                   FROM pedidos WHERE DATE(fecha_pedido) BETWEEN %s AND %s AND estado != 'cancelado'"""
        params = [inicio.strftime('%Y-%m-%d'), fin.strftime('%Y-%m-%d')]
        if tienda_id:
            query += ' AND tienda_id = %s'
            params.append(tienda_id)
        return db.execute(query, params).fetchone()

    actual = get_stats(inicio_mes_actual, hoy)
    anterior = get_stats(inicio_mes_anterior, fin_mes_anterior)
    db.close()

    def cambio(a, b):
        if b == 0: return 100 if a > 0 else 0
        return round((a - b) / b * 100, 1)

    return jsonify({
        'mes_actual': {
            'nombre': hoy.strftime('%B %Y'),
            'pedidos': actual['pedidos'],
            'ventas': float(actual['ventas']),
            'clientes': actual['clientes'],
            'ticket_promedio': float(actual['ventas'] / actual['pedidos']) if actual['pedidos'] > 0 else 0
        },
        'mes_anterior': {
            'nombre': fin_mes_anterior.strftime('%B %Y'),
            'pedidos': anterior['pedidos'],
            'ventas': float(anterior['ventas']),
            'clientes': anterior['clientes'],
            'ticket_promedio': float(anterior['ventas'] / anterior['pedidos']) if anterior['pedidos'] > 0 else 0
        },
        'cambio': {
            'pedidos': cambio(actual['pedidos'], anterior['pedidos']),
            'ventas': cambio(float(actual['ventas']), float(anterior['ventas'])),
            'clientes': cambio(actual['clientes'], anterior['clientes'])
        }
    })


@superadmin_bp.route('/reportes/api/ventas-por-dia-semana')
@requiere_superadmin
def api_ventas_por_dia_semana():
    """Ventas agrupadas por dia de la semana"""
    db = get_connection()
    fecha_desde = request.args.get('desde', (datetime.now() - timedelta(days=30)).strftime('%Y-%m-%d'))
    fecha_hasta = request.args.get('hasta', datetime.now().strftime('%Y-%m-%d'))
    tienda_id = request.args.get('tienda')

    query = """SELECT DAYOFWEEK(fecha_pedido) as dia, COUNT(*) as pedidos, SUM(total) as ventas
               FROM pedidos WHERE DATE(fecha_pedido) BETWEEN %s AND %s AND estado != 'cancelado'"""
    params = [fecha_desde, fecha_hasta]
    if tienda_id:
        query += ' AND tienda_id = %s'
        params.append(tienda_id)
    query += ' GROUP BY DAYOFWEEK(fecha_pedido) ORDER BY dia'

    resultados = db.execute(query, params).fetchall()
    db.close()

    dias_nombres = ['Domingo', 'Lunes', 'Martes', 'Miercoles', 'Jueves', 'Viernes', 'Sabado']
    dias_dict = {r['dia']: r for r in resultados}

    dias = []
    for d in range(1, 8):
        if d in dias_dict:
            dias.append({'dia': d, 'nombre': dias_nombres[d-1], 'pedidos': dias_dict[d]['pedidos'], 'ventas': float(dias_dict[d]['ventas'])})
        else:
            dias.append({'dia': d, 'nombre': dias_nombres[d-1], 'pedidos': 0, 'ventas': 0})

    return jsonify(dias)


@superadmin_bp.route('/reportes/api/top-clientes')
@requiere_superadmin
def api_top_clientes():
    """Top clientes por ventas"""
    db = get_connection()
    fecha_desde = request.args.get('desde', (datetime.now() - timedelta(days=30)).strftime('%Y-%m-%d'))
    fecha_hasta = request.args.get('hasta', datetime.now().strftime('%Y-%m-%d'))
    tienda_id = request.args.get('tienda')
    limite = request.args.get('limite', 20, type=int)

    query = """SELECT c.nombre, c.telefono, c.email, COUNT(p.id) as pedidos,
               SUM(p.total) as total, MAX(p.fecha_pedido) as ultima_compra
               FROM clientes c JOIN pedidos p ON c.id = p.cliente_id
               WHERE DATE(p.fecha_pedido) BETWEEN %s AND %s AND p.estado != 'cancelado'"""
    params = [fecha_desde, fecha_hasta]
    if tienda_id:
        query += ' AND p.tienda_id = %s'
        params.append(tienda_id)
    query += f' GROUP BY c.id ORDER BY total DESC LIMIT {limite}'

    clientes = db.execute(query, params).fetchall()
    db.close()

    resultado = []
    for c in clientes:
        resultado.append({
            'nombre': c['nombre'],
            'telefono': c['telefono'],
            'email': c['email'],
            'pedidos': c['pedidos'],
            'total': float(c['total']),
            'ultima_compra': str(c['ultima_compra'])[:16] if c['ultima_compra'] else None
        })

    return jsonify(resultado)


# ============ CAMBIAR CONTRASEÑA DE USUARIO DE TIENDA ============

@superadmin_bp.route('/api/tiendas/<int:tienda_id>/usuarios')
@requiere_superadmin
def api_usuarios_tienda(tienda_id):
    """Obtener usuarios de una tienda"""
    db = get_connection()
    usuarios = db.execute("""
        SELECT id, email, nombre, rol, activo
        FROM usuarios WHERE tienda_id = %s ORDER BY rol
    """, (tienda_id,)).fetchall()
    db.close()

    return jsonify([dict(u) for u in usuarios])


@superadmin_bp.route('/api/usuarios/<int:usuario_id>/cambiar-password', methods=['POST'])
@requiere_superadmin
def api_cambiar_password_usuario(usuario_id):
    """Cambiar contraseña de un usuario"""
    from werkzeug.security import generate_password_hash

    data = request.get_json()
    nueva_password = data.get('password')

    if not nueva_password or len(nueva_password) < 4:
        return jsonify({'success': False, 'error': 'La contraseña debe tener al menos 4 caracteres'}), 400

    db = get_connection()

    # Verificar que el usuario existe
    usuario = db.execute("SELECT id, email, nombre FROM usuarios WHERE id = %s", (usuario_id,)).fetchone()
    if not usuario:
        db.close()
        return jsonify({'success': False, 'error': 'Usuario no encontrado'}), 404

    # Actualizar contraseña
    db.execute("""
        UPDATE usuarios SET password_hash = %s WHERE id = %s
    """, (generate_password_hash(nueva_password), usuario_id))
    db.close()

    return jsonify({
        'success': True,
        'message': f'Contraseña actualizada para {usuario["nombre"]} ({usuario["email"]})'
    })


@superadmin_bp.route('/api/tiendas/<int:tienda_id>/reset-passwords', methods=['POST'])
@requiere_superadmin
def api_reset_passwords_tienda(tienda_id):
    """Resetear todas las contraseñas de una tienda y devolver las nuevas"""
    from werkzeug.security import generate_password_hash
    import random
    import string

    def gen_password():
        return ''.join(random.choices(string.ascii_letters + string.digits, k=8))

    db = get_connection()

    # Obtener usuarios de la tienda
    usuarios = db.execute("""
        SELECT id, email, nombre, rol FROM usuarios WHERE tienda_id = %s
    """, (tienda_id,)).fetchall()

    if not usuarios:
        db.close()
        return jsonify({'success': False, 'error': 'No hay usuarios en esta tienda'}), 404

    nuevas_credenciales = []

    for usuario in usuarios:
        nueva_password = gen_password()
        db.execute("""
            UPDATE usuarios SET password_hash = %s WHERE id = %s
        """, (generate_password_hash(nueva_password), usuario['id']))

        nuevas_credenciales.append({
            'email': usuario['email'],
            'nombre': usuario['nombre'],
            'rol': usuario['rol'],
            'password': nueva_password
        })

    db.close()

    return jsonify({
        'success': True,
        'message': f'Se resetearon {len(nuevas_credenciales)} contraseñas',
        'credenciales': nuevas_credenciales
    })


# ============ RUTAS DE BACKUPS ============
import subprocess
from flask import send_file

@superadmin_bp.route('/backups')
@requiere_superadmin
def backups():
    """Página de gestión de backups"""
    return render_template('superadmin/backups.html')

@superadmin_bp.route('/api/backups', methods=['GET'])
@requiere_superadmin
def api_backups_listar():
    """Listar backups disponibles"""
    import os
    from datetime import datetime

    backup_dir = '/app/data/backups'
    backups = []

    try:
        if os.path.exists(backup_dir):
            for filename in os.listdir(backup_dir):
                if filename.endswith('.sql.gz'):
                    filepath = os.path.join(backup_dir, filename)
                    stat = os.stat(filepath)

                    # Extraer fecha del nombre del archivo
                    # Formato: restaurantes_20251214_014542.sql.gz
                    try:
                        date_part = filename.replace('restaurantes_', '').replace('.sql.gz', '')
                        fecha = datetime.strptime(date_part, '%Y%m%d_%H%M%S')
                        fecha_str = fecha.strftime('%d/%m/%Y %H:%M:%S')
                    except:
                        fecha_str = datetime.fromtimestamp(stat.st_mtime).strftime('%d/%m/%Y %H:%M:%S')

                    # Tamaño legible
                    size_bytes = stat.st_size
                    if size_bytes < 1024:
                        size_str = f"{size_bytes} B"
                    elif size_bytes < 1024 * 1024:
                        size_str = f"{size_bytes / 1024:.1f} KB"
                    else:
                        size_str = f"{size_bytes / (1024 * 1024):.1f} MB"

                    backups.append({
                        'nombre': filename,
                        'fecha': fecha_str,
                        'tamano': size_str,
                        'tamano_bytes': size_bytes
                    })

            # Ordenar por fecha (más reciente primero)
            backups.sort(key=lambda x: x['nombre'], reverse=True)
    except Exception as e:
        return jsonify({'error': str(e)}), 500

    return jsonify({'backups': backups})

@superadmin_bp.route('/api/backups/crear', methods=['POST'])
@requiere_superadmin
@requiere_csrf
def api_backup_crear():
    """Crear un nuevo backup manualmente usando PyMySQL"""
    import gzip
    from datetime import datetime
    import os
    import time

    try:
        backup_dir = '/app/data/backups'
        os.makedirs(backup_dir, exist_ok=True)

        date_str = datetime.now().strftime('%Y%m%d_%H%M%S')
        gz_file = f'{backup_dir}/restaurantes_{date_str}.sql.gz'

        db = get_connection()

        # Obtener todas las tablas
        tablas = db.execute("SHOW TABLES").fetchall()

        sql_content = []
        sql_content.append(f"-- Backup generado: {datetime.now().isoformat()}")
        sql_content.append("SET FOREIGN_KEY_CHECKS=0;")
        sql_content.append("")

        for (tabla,) in tablas:
            # CREATE TABLE
            create = db.execute(f"SHOW CREATE TABLE `{tabla}`").fetchone()
            sql_content.append(f"DROP TABLE IF EXISTS `{tabla}`;")
            sql_content.append(create[1] + ";")
            sql_content.append("")

            # INSERT DATA
            rows = db.execute(f"SELECT * FROM `{tabla}`").fetchall()
            if rows:
                # Obtener nombres de columnas
                cols = db.execute(f"SHOW COLUMNS FROM `{tabla}`").fetchall()
                col_names = [c[0] for c in cols]

                for row in rows:
                    values = []
                    for val in row:
                        if val is None:
                            values.append("NULL")
                        elif isinstance(val, (int, float)):
                            values.append(str(val))
                        elif isinstance(val, datetime):
                            values.append(f"'{val.strftime('%Y-%m-%d %H:%M:%S')}'")
                        else:
                            # Escapar comillas
                            escaped = str(val).replace("\\", "\\\\").replace("'", "\\'")
                            values.append(f"'{escaped}'")

                    sql_content.append(f"INSERT INTO `{tabla}` VALUES ({', '.join(values)});")

            sql_content.append("")

        sql_content.append("SET FOREIGN_KEY_CHECKS=1;")

        # Comprimir y guardar
        with gzip.open(gz_file, 'wt', encoding='utf-8') as f:
            f.write("\n".join(sql_content))

        # Limpiar backups antiguos (más de 7 días)
        now = time.time()
        for filename in os.listdir(backup_dir):
            if filename.endswith('.sql.gz'):
                filepath = os.path.join(backup_dir, filename)
                if os.stat(filepath).st_mtime < now - 7 * 86400:
                    os.remove(filepath)

        return jsonify({'success': True, 'mensaje': f'Backup creado: restaurantes_{date_str}.sql.gz'})

    except Exception as e:
        import traceback
        return jsonify({'error': str(e), 'trace': traceback.format_exc()}), 500

@superadmin_bp.route('/api/backups/descargar/<nombre>')
@requiere_superadmin
def api_backup_descargar(nombre):
    """Descargar un backup específico"""
    import re
    import os

    # Validar nombre del archivo (seguridad)
    if not re.match(r'^restaurantes_\d{8}_\d{6}\.sql\.gz$', nombre):
        return jsonify({'error': 'Nombre de archivo inválido'}), 400

    filepath = f'/app/data/backups/{nombre}'

    if not os.path.exists(filepath):
        return jsonify({'error': 'Archivo no encontrado'}), 404

    return send_file(
        filepath,
        as_attachment=True,
        download_name=nombre
    )

@superadmin_bp.route('/api/backups/eliminar/<nombre>', methods=['DELETE'])
@requiere_superadmin
@requiere_csrf
def api_backup_eliminar(nombre):
    """Eliminar un backup específico"""
    import re
    import os

    # Validar nombre del archivo (seguridad)
    if not re.match(r'^restaurantes_\d{8}_\d{6}\.sql\.gz$', nombre):
        return jsonify({'error': 'Nombre de archivo inválido'}), 400

    filepath = f'/app/data/backups/{nombre}'

    if not os.path.exists(filepath):
        return jsonify({'error': 'Archivo no encontrado'}), 404

    try:
        os.remove(filepath)
        return jsonify({'success': True, 'mensaje': 'Backup eliminado'})
    except Exception as e:
        return jsonify({'error': str(e)}), 500
